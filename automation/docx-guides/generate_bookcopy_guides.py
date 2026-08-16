from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('Documents/Student-Guides/BookCopy')
OUT.mkdir(parents=True, exist_ok=True)
CONTRACT = 'LIBRARY-FULLSTACK-V1'

TRACKS = {
'T31': dict(student='Student 31', op='Add Book Copy', code='31', service='CreateBookCopyService', impl='CreateBookCopyServiceImpl', endpoint='POST /rest/book-copies', success='Book Copy Added Successfully', ui='Update Book Copy UI', unit='Call createBookCopy(null) after the successful Add test works.', integration='The database already contains accession number ACC-0001. Try to add another Book Copy with the same persisted accession number.'),
'T32': dict(student='Student 32', op='Read Book Copy', code='32', service='ReadBookCopyService', impl='ReadBookCopyServiceImpl', endpoint='GET /rest/book-copies/{id}', success='Book Copy Read Successfully', ui='Withdraw Book Copy UI', unit='Request Book Copy ID 999 after the successful Read test works.', integration='Book Copy ID 3 physically exists as ACC-0003 but has status WITHDRAWN. Normal Read must not expose this non-current copy.'),
'T33': dict(student='Student 33', op='Update Book Copy', code='33', service='UpdateBookCopyService', impl='UpdateBookCopyServiceImpl', endpoint='PUT /rest/book-copies/{id}', success='Book Copy Updated Successfully', ui='Search Book Copy UI', unit='Call updateBookCopy(4L, null) after the successful Update test works.', integration='Book Copy ID 1 owns ACC-0001 and ID 4 owns ACC-0004. Try to update ID 4 to ACC-0001. Exclude the current row from duplicate detection.'),
'T34': dict(student='Student 34', op='Withdraw Book Copy', code='34', service='DeleteBookCopyService', impl='DeleteBookCopyServiceImpl', endpoint='DELETE /rest/book-copies/{id}', success='Book Copy Withdrawn Successfully', ui='Add Book Copy UI', unit='Request Book Copy ID 999 after the normal withdrawal path works.', integration='Book Copy ID 4 has no ACTIVE Book Issue and can be withdrawn. Book Copy ID 2 is used by ACTIVE issue ISS-0001 and must not be withdrawn.'),
'T35': dict(student='Student 35', op='Search Book Copy', code='35', service='SearchBookCopyService', impl='SearchBookCopyServiceImpl', endpoint='GET /rest/book-copies/search?text={text}', success='Book Copy Search Completed Successfully', ui='Book Copy List / Search Results UI', unit='Call searchBookCopy(null) after normal Search works. The first implementation calls trim().', integration='Search for acc-0001 while PostgreSQL stores ACC-0001. The completed search must be current-only and case-independent.'),
}

def shade(cell, fill):
    p=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),fill); p.append(sh)

def setup(doc):
    for s in doc.sections:
        s.top_margin=Inches(.62); s.bottom_margin=Inches(.62); s.left_margin=Inches(.72); s.right_margin=Inches(.72)
    doc.styles['Normal'].font.name='Aptos'; doc.styles['Normal'].font.size=Pt(10.2)
    for n,sz in [('Title',21),('Heading 1',15),('Heading 2',12.7),('Heading 3',11.1)]:
        doc.styles[n].font.name='Aptos'; doc.styles[n].font.size=Pt(sz)

def title(doc,a,b):
    p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(a)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(b); r.bold=True

def table(doc, rows, headers=('Item','Meaning / Value')):
    t=doc.add_table(rows=1,cols=2); t.style='Table Grid'; t.cell(0,0).text=headers[0]; t.cell(0,1).text=headers[1]; shade(t.cell(0,0),'D9EAF7'); shade(t.cell(0,1),'D9EAF7')
    for a,b in rows:
        c=t.add_row().cells; c[0].text=str(a); c[1].text=str(b)

def code(doc,s):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.12); r=p.add_run(s.strip()); r.font.name='Consolas'; r.font.size=Pt(8.05); sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),'F3F4F6'); p._p.get_or_add_pPr().append(sh)

def note(doc,s):
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'; shade(t.cell(0,0),'FFF7D6'); p=t.cell(0,0).paragraphs[0]; r=p.add_run('Important: '); r.bold=True; p.add_run(s)

def bullets(doc, items):
    for x in items: doc.add_paragraph(x,style='List Bullet')

def sig(t):
    return {'T31':'BookCopyResponseDto createBookCopy(BookCopyCreateRequestDto request)','T32':'BookCopyResponseDto readBookCopy(Long id)','T33':'BookCopyResponseDto updateBookCopy(Long id, BookCopyUpdateRequestDto request)','T34':'BookCopyResponseDto deleteBookCopy(Long id)','T35':'List<BookCopyResponseDto> searchBookCopy(String text)'}[t]

def normal(t):
    return {
'T31':'''@Autowired private BookCopyDao bookCopyDao;\n@Autowired private BookCopyDtoDoMapper mapper;\n\n@Override\npublic BookCopyResponseDto createBookCopy(BookCopyCreateRequestDto request) {\n    BookCopyDO copy = mapper.toDO(request);\n    return mapper.toResponse(bookCopyDao.save(copy));\n}''',
'T32':'''@Autowired private BookCopyDao bookCopyDao;\n@Autowired private BookCopyDtoDoMapper mapper;\n\n@Override\npublic BookCopyResponseDto readBookCopy(Long id) {\n    BookCopyDO copy = bookCopyDao.findById(id).orElse(null);\n    return mapper.toResponse(copy);\n}''',
'T33':'''@Autowired private BookCopyDao bookCopyDao;\n@Autowired private BookCopyDtoDoMapper mapper;\n\n@Override\npublic BookCopyResponseDto updateBookCopy(Long id, BookCopyUpdateRequestDto request) {\n    BookCopyDO copy = bookCopyDao.findById(id).orElse(null);\n    mapper.applyUpdate(request, copy);\n    return mapper.toResponse(bookCopyDao.save(copy));\n}''',
'T34':'''@Autowired private BookCopyDao bookCopyDao;\n@Autowired private BookCopyDtoDoMapper mapper;\n\n@Override\npublic BookCopyResponseDto deleteBookCopy(Long id) {\n    BookCopyDO copy = bookCopyDao.findById(id).orElse(null);\n    copy.setStatus("WITHDRAWN");\n    return mapper.toResponse(bookCopyDao.save(copy));\n}''',
'T35':'''@Autowired private BookCopyDao bookCopyDao;\n@Autowired private BookCopyDtoDoMapper mapper;\n\n@Override\npublic List<BookCopyResponseDto> searchBookCopy(String text) {\n    String normalized = text.trim();\n    return bookCopyDao.search(normalized).stream().map(mapper::toResponse).toList();\n}'''}[t]

def corrected(t):
    return {
'T31':'''@Override\npublic BookCopyResponseDto createBookCopy(BookCopyCreateRequestDto request) {\n    if (request == null) throw new IllegalArgumentException("Book Copy request is required");\n    String accession = request.getAccessionNumber().trim();\n    bookCopyDao.findByNormalizedBusinessKey(accession).ifPresent(existing -> { throw new IllegalStateException("Accession number already exists"); });\n    request.setAccessionNumber(accession);\n    return mapper.toResponse(bookCopyDao.save(mapper.toDO(request)));\n}''',
'T32':'''@Override\npublic BookCopyResponseDto readBookCopy(Long id) {\n    if (id == null) throw new IllegalArgumentException("Book Copy ID is required");\n    BookCopyDO copy = bookCopyDao.findCurrentById(id).orElseThrow(() -> new IllegalArgumentException("Book Copy not found"));\n    return mapper.toResponse(copy);\n}''',
'T33':'''@Override\npublic BookCopyResponseDto updateBookCopy(Long id, BookCopyUpdateRequestDto request) {\n    if (id == null || request == null) throw new IllegalArgumentException("Book Copy ID and request are required");\n    BookCopyDO copy = bookCopyDao.findCurrentById(id).orElseThrow(() -> new IllegalArgumentException("Book Copy not found"));\n    String accession = request.getAccessionNumber().trim();\n    bookCopyDao.findByNormalizedBusinessKey(accession).filter(other -> !other.getBookCopyId().equals(id)).ifPresent(other -> { throw new IllegalStateException("Accession number already exists"); });\n    request.setAccessionNumber(accession); mapper.applyUpdate(request, copy);\n    return mapper.toResponse(bookCopyDao.save(copy));\n}''',
'T34':'''@Override\npublic BookCopyResponseDto deleteBookCopy(Long id) {\n    if (id == null) throw new IllegalArgumentException("Book Copy ID is required");\n    BookCopyDO copy = bookCopyDao.findCurrentById(id).orElseThrow(() -> new IllegalArgumentException("Book Copy not found"));\n    if (bookCopyDao.countDependencies(id) > 0) throw new IllegalStateException("Book Copy has an active Book Issue");\n    copy.setStatus("WITHDRAWN");\n    return mapper.toResponse(bookCopyDao.save(copy));\n}''',
'T35':'''@Override\npublic List<BookCopyResponseDto> searchBookCopy(String text) {\n    if (text == null) throw new IllegalArgumentException("Search text is required");\n    String normalized = text.trim();\n    if (normalized.isBlank()) throw new IllegalArgumentException("Search text is required");\n    return bookCopyDao.searchIgnoreCase(normalized).stream().map(mapper::toResponse).toList();\n}'''}[t]

def observe(t):
    return {'T31':'service.createBookCopy(null);','T32':'when(bookCopyDao.findById(999L)).thenReturn(Optional.empty());\nservice.readBookCopy(999L);','T33':'when(bookCopyDao.findById(4L)).thenReturn(Optional.of(existing));\nservice.updateBookCopy(4L, null);','T34':'when(bookCopyDao.findById(999L)).thenReturn(Optional.empty());\nservice.deleteBookCopy(999L);','T35':'service.searchBookCopy(null);'}[t]

def itest(t):
    return {
'T31':'''@SpringBootTest\nclass CreateBookCopyLocalIntegrationTest {\n @Autowired CreateBookCopyService service;\n @Test void persistedDuplicateAccessionNumberIsRejected() {\n  assertThrows(IllegalStateException.class, () -> service.createBookCopy(request("ACC-0001",1L,"AVAILABLE")));\n }\n}''',
'T32':'''@SpringBootTest\nclass ReadBookCopyLocalIntegrationTest {\n @Autowired ReadBookCopyService service;\n @Test void currentCopyIsRead() { assertEquals("ACC-0004", service.readBookCopy(4L).getAccessionNumber()); }\n @Test void withdrawnCopyIsNotExposed() { assertThrows(IllegalArgumentException.class, () -> service.readBookCopy(3L)); }\n}''',
'T33':'''@SpringBootTest\nclass UpdateBookCopyLocalIntegrationTest {\n @Autowired UpdateBookCopyService service;\n @Test void accessionOwnedByAnotherCopyIsRejected() {\n  assertThrows(IllegalStateException.class, () -> service.updateBookCopy(4L, request("ACC-0001",4L,"AVAILABLE")));\n }\n}''',
'T34':'''@SpringBootTest\nclass WithdrawBookCopyLocalIntegrationTest {\n @Autowired DeleteBookCopyService service;\n @Test void copyWithoutActiveIssueCanBeWithdrawn() { assertEquals("WITHDRAWN", service.deleteBookCopy(4L).getStatus()); }\n @Test void copyWithActiveIssueCannotBeWithdrawn() { assertThrows(IllegalStateException.class, () -> service.deleteBookCopy(2L)); }\n}''',
'T35':'''@SpringBootTest\nclass SearchBookCopyLocalIntegrationTest {\n @Autowired SearchBookCopyService service;\n @Test void lowercaseSearchFindsUppercaseAccession() { assertTrue(service.searchBookCopy("acc-0001").stream().anyMatch(c -> "ACC-0001".equals(c.getAccessionNumber()))); }\n @Test void validNoMatchIsEmptyList() { assertTrue(service.searchBookCopy("NO-SUCH-COPY").isEmpty()); }\n}'''}[t]

def initial(t,d):
    doc=Document(); setup(doc); title(doc,f'{t} - {d["op"]} - Initial API Contract',f'{d["student"]} | {CONTRACT} | Book Copy Module')
    doc.add_heading('1. Fixed Assignment',1); table(doc,[('Track',t),('Student',d['student']),('Operation',d['op']),('Service Code',d['code']),('Endpoint',d['endpoint']),('Service Interface',d['service']),('Method',sig(t)),('Success responseCode','00'),('Success Message',d['success'])])
    doc.add_heading('2. JSON Response Envelope',1); code(doc,f'{{\n "serviceCode":"{d["code"]}",\n "responseCode":"00",\n "message":"{d["success"]}",\n "data":{{}}\n}}')
    doc.add_heading('3. Book Copy Fields',1); table(doc,[('bookCopyId','Database identifier; response only.'),('accessionNumber','Business key such as ACC-0001.'),('bookId','Parent Book identifier.'),('status','AVAILABLE, ISSUED or WITHDRAWN.')])
    doc.add_heading('4. Training Data',1); table(doc,[('Copy 1','ACC-0001 / Book 1 / AVAILABLE'),('Copy 2','ACC-0002 / Book 1 / ISSUED; used by ACTIVE issue ISS-0001'),('Copy 3','ACC-0003 / Book 2 / WITHDRAWN'),('Copy 4','ACC-0004 / Book 4 / AVAILABLE; no active issue')])
    doc.add_heading('5. Ownership',1); table(doc,[('Presenter','REST Controller, DTO/DO, Mapper, base DAO, Flyway, JSON envelope and Thymeleaf framework.'),('Student',f'{d["impl"]}, Unit/Integration tests and {d["ui"]}.')]); note(doc,'WITHDRAW is a logical state change. The row remains in the database with status WITHDRAWN.')
    return doc

def progressive(t,d):
    doc=Document(); setup(doc); title(doc,f'{d["student"]} - {d["op"]} Progressive Testing Guide',f'{t} | Book Copy Module | {CONTRACT}')
    doc.add_paragraph('Follow the accepted progressive sequence: empty service, direct Unit Test invocation, normal successful implementation, successful Unit Test, next realistic condition, observation, correction, JaCoCo, local PostgreSQL, complete Integration Test, Testcontainers, regression and frontend integration.')
    doc.add_heading('1. Assignment and Application Flow',1); table(doc,[('Backend Assignment',d['op']),('Service',d['service']),('Implementation',d['impl']),('REST Endpoint',d['endpoint']),('Service Code',d['code']),('Frontend Assignment',d['ui']),('REST Controller','Presenter-owned')]); table(doc,[('Browser / client','Sends HTTP.'),('Spring + Jackson','Converts JSON body to supplied DTO.'),('REST Controller','Calls fixed service.'),('Service','Student behavior.'),('Mapper','DTO <-> DO.'),('DAO','JPA persistence.'),('PostgreSQL','Real data.')]); note(doc,'The service works with Java DTOs, not raw JSON.')
    doc.add_heading('2. Create and Invoke the Empty Service',1); code(doc,f'@Service\npublic class {d["impl"]} implements {d["service"]} {{\n @Override public {sig(t)} {{ ' + ('return java.util.Collections.emptyList();' if t=='T35' else 'return null;') + ' }\n}'); code(doc,f'@ExtendWith(MockitoExtension.class)\nclass {d["impl"]}Test {{\n @InjectMocks private {d["impl"]} service;\n @Test void shouldInvokeService() {{ /* invoke the assigned method */ }}\n}}'); doc.add_paragraph('Run this focused method from Eclipse with Run As -> JUnit Test. The empty service proves the component can be called independently before the full Spring chain is involved.')
    doc.add_heading('3. Implement the Normal Successful Path',1); code(doc,normal(t)); doc.add_paragraph('Add a Mockito Unit Test for one valid path and keep it as a permanent regression test.')
    doc.add_heading('4. Try the Next Realistic Condition',1); doc.add_paragraph(d['unit']); doc.add_paragraph('Do not change the service first. Run the test and observe the current result.'); code(doc,observe(t)); doc.add_heading('4.1 Understand NullPointerException',2); doc.add_paragraph('A NullPointerException occurs when code calls a method or accesses state through a null reference. Read the first stack-trace line belonging to the service. Handle the application condition before dereferencing the null value.')
    doc.add_heading('5. Improve the Service',1); code(doc,corrected(t)); doc.add_paragraph('Rerun the new test and the original successful Unit Test. The correction must preserve the previously working path.')
    doc.add_heading('6. JaCoCo Coverage',1); code(doc,'mvnw.cmd clean test'); table(doc,[('Green','Executed.'),('Yellow','Partially covered branch.'),('Red','Not executed; room for uncertainty.')]); doc.add_paragraph('Open target/site/jacoco/index.html and inspect the Book Copy service.')
    doc.add_heading('7. Why Integration Testing is Required',1); doc.add_paragraph('Mockito cannot prove real accession-number duplication, WITHDRAWN state filtering, active Book Issue dependencies or PostgreSQL case behavior. Integration Testing joins the real Mapper, DAO and database.'); table(doc,[('Unit Test','Service + mocks.'),('Local Integration','Spring + real Mapper + DAO + PostgreSQL.'),('Testcontainers','Same path with temporary PostgreSQL.')])
    doc.add_heading('8. PostgreSQL JDBC and Commented MySQL Reference',1); code(doc,'spring.datasource.url=jdbc:postgresql://localhost:5432/library_full_stack\nspring.datasource.username=postgres\nspring.datasource.password=postgres\nspring.datasource.driver-class-name=org.postgresql.Driver'); code(doc,'<dependency>\n <groupId>org.postgresql</groupId>\n <artifactId>postgresql</artifactId>\n <scope>runtime</scope>\n</dependency>\n\n<!-- MySQL reference only\n<dependency>\n <groupId>com.mysql</groupId>\n <artifactId>mysql-connector-j</artifactId>\n <scope>runtime</scope>\n</dependency>\n-->'); code(doc,'# MySQL reference only - keep commented\n# spring.datasource.driver-class-name=com.mysql.cj.jdbc.Driver')
    doc.add_heading('9. Inspect Real Data with DBeaver',1); code(doc,'SELECT book_copy_id, accession_number, fk_book, status FROM tbl_book_copy ORDER BY book_copy_id;\n\nSELECT book_issue_id, issue_number, fk_book_copy, status FROM tbl_book_issue ORDER BY book_issue_id;'); doc.add_paragraph(d['integration'])
    doc.add_heading('10. Complete Local Integration Test',1); code(doc,itest(t))
    if t=='T35': doc.add_heading('10.1 Corrected PostgreSQL Search',2); code(doc,'@Transactional(readOnly = true)\npublic List<BookCopyDO> searchIgnoreCase(String text) {\n String pattern = "%" + text.toLowerCase() + "%";\n return entityManager.createQuery("select c from BookCopyDO c where c.status <> \'WITHDRAWN\' and lower(c.accessionNumber) like :pattern order by c.bookCopyId", BookCopyDO.class).setParameter("pattern",pattern).getResultList();\n}')
    doc.add_heading('11. PostgreSQL Testcontainers',1); code(doc,'@Container\n@ServiceConnection\nstatic PostgreSQLContainer postgres = new PostgreSQLContainer(DockerImageName.parse("postgres:18"));'); doc.add_paragraph('Add @Testcontainers to the Integration Test class and repeat the same test methods. Flyway recreates the deterministic training data.')
    doc.add_heading('12. Final Regression',1); code(doc,'mvnw.cmd clean test'); bullets(doc,['Focused Unit Test green.','Original successful Unit Test still green.','Local PostgreSQL Integration Test green.','PostgreSQL Testcontainers test green.','JaCoCo reviewed.'])
    doc.add_heading('13. Student-owned Files',1); table(doc,[('Service',d['impl']+'.java'),('Tests','Track-specific Unit and Integration Test classes.'),('Frontend',d['ui']),('Presenter-owned','REST Controller, DTO/DO, Mapper, base DAO, Flyway and Thymeleaf framework.')]);
    if t=='T35': note(doc,'T35 owns only the case-independent BookCopyDao search method named in this guide.')
    doc.add_heading('14. Frontend Assignment',1); doc.add_paragraph(f'{d["student"]} implements the {d["ui"]}. Prove main.js and the button handler, build one working JavaScript file, test against Presenter STUB/support, then refactor and integrate the partner service without changing the /rest URL.'); code(doc,'<script type="module" src="/src/main.js"></script>\n\nconsole.log("main.js loaded");\ndocument.querySelector("#actionButton").addEventListener("click", () => console.log("button reached JavaScript"));')
    doc.add_heading('15. HTTP Transport Handling',1); code(doc,'if (response.status === 404) { showMessage("NOT Found"); return; }\nif (!response.ok) { showMessage("Error Encountered, Please contact Administrator"); return; }\nconst body = await response.json();\nshowMessage(body.message);')
    doc.add_heading('16. Refactor After the Page Works',1); table(doc,[('src/main.js','Starts the page.'),('src/api/bookCopyApi.js','Contains REST calls.'),('src/forms/bookCopyForm.js','Reads form values.'),('src/views/bookCopyView.js','Renders records/messages.')]); bullets(doc,['Presenter STUB/support works.','Partner backend works with unchanged URL.','Assigned files committed and pushed.'])
    return doc

def updated(t,d):
    doc=Document(); setup(doc); title(doc,f'{t} - {d["op"]} - Updated API Contract',f'{CONTRACT} | Book Copy Module'); doc.add_heading('1. Fixed Endpoint',1); table(doc,[('Track',t),('Service Code',d['code']),('Endpoint',d['endpoint']),('Controlled application result','HTTP 200')]); doc.add_heading('2. Response Envelope',1); code(doc,'{\n "serviceCode":"XX",\n "responseCode":"00",\n "message":"...",\n "data":{}\n}')
    rows=[('00',d['success'])]
    if t=='T31': rows += [('01','Invalid / missing Book Copy request'),('03','Accession number already exists')]
    elif t=='T32': rows += [('01','Book Copy ID is required'),('02','Book Copy not found / WITHDRAWN copy not exposed')]
    elif t=='T33': rows += [('01','Invalid Book Copy ID/request'),('02','Book Copy not found'),('03','Accession number belongs to another Book Copy')]
    elif t=='T34': rows += [('01','Book Copy ID is required'),('02','Book Copy not found'),('05','Book Copy has an ACTIVE Book Issue')]
    else: rows += [('01','Search text is required or blank')]
    doc.add_heading('3. Response Codes',1); table(doc,rows,('responseCode','Meaning')); doc.add_heading('4. Transport Errors',1); table(doc,[('HTTP 404','NOT Found'),('Other non-200','Error Encountered, Please contact Administrator')]); note(doc,'A valid Search with zero matches returns responseCode 00 and an empty list.' if t=='T35' else 'Controlled application results use the frozen JSON response envelope and HTTP 200.'); return doc

for t,d in TRACKS.items():
    folder=OUT/f'{t}_{d["op"].replace(" ","_")}'
    folder.mkdir(parents=True,exist_ok=True)
    initial(t,d).save(folder/f'{t}_01_Initial_API_Contract.docx')
    progressive(t,d).save(folder/f'{t}_02_Progressive_Development_Guide.docx')
    updated(t,d).save(folder/f'{t}_03_Updated_API_Contract.docx')

print('Generated Book Copy guides under', OUT)
