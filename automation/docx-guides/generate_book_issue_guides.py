from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import textwrap

CONTRACT = "LIBRARY-FULLSTACK-V1"
OUT = Path("Documents/Student-Guides/Book-Issue")

TRACKS = {
    "T41": dict(student="Student 41", operation="Create Book Issue", code="41", service="CreateBookIssueService", impl="CreateBookIssueServiceImpl", endpoint="POST /rest/issues", success="Book Issue Created Successfully", frontend="Update Book Issue UI", unit="After the valid Create test is green, call createBookIssue(null). Run the test before adding validation.", integration="The training database already contains issueNumber ISS-0001. Try to create another Book Issue using the same persisted issue number."),
    "T42": dict(student="Student 42", operation="Read Book Issue", code="42", service="ReadBookIssueService", impl="ReadBookIssueServiceImpl", endpoint="GET /rest/issues/{id}", success="Book Issue Read Successfully", frontend="Cancel Book Issue UI", unit="After the valid Read test is green, request Book Issue ID 999. Configure BookIssueDao to return Optional.empty() and observe the current behavior.", integration="Book Issue ID 2 physically exists as ISS-0002 but its status is RETURNED. Normal Read must not expose this historical Issue as a current ACTIVE Issue."),
    "T43": dict(student="Student 43", operation="Update Book Issue", code="43", service="UpdateBookIssueService", impl="UpdateBookIssueServiceImpl", endpoint="PUT /rest/issues/{id}", success="Book Issue Updated Successfully", frontend="Search Book Issue UI", unit="After the valid Update test is green, call updateBookIssue(1L, null). Run the test before adding validation.", integration="Book Issue ID 1 owns ISS-0001 and ID 2 owns ISS-0002. Try to update ID 2 to issueNumber ISS-0001. The current row must be excluded from duplicate detection."),
    "T44": dict(student="Student 44", operation="Cancel Book Issue", code="44", service="DeleteBookIssueService", impl="DeleteBookIssueServiceImpl", endpoint="DELETE /rest/issues/{id}", success="Book Issue Cancelled Successfully", frontend="Create Book Issue UI", unit="After the normal Cancel path is green, request Book Issue ID 999. Run the test before adding not-found handling.", integration="Training Issue ISS-0005 remains ACTIVE but has a persisted non-VOID Return record RET-0003. It represents a return already recorded while the Issue state update is still pending. The Issue must not be cancelled once that Return exists."),
    "T45": dict(student="Student 45", operation="Search Book Issue", code="45", service="SearchBookIssueService", impl="SearchBookIssueServiceImpl", endpoint="GET /rest/issues/search?text={text}", success="Book Issue Search Completed Successfully", frontend="Book Issue List / Search Results UI", unit="After normal Search is green, call searchBookIssue(null). The normal first implementation calls trim(), so run the test before adding validation.", integration="Search for iss-0001 while PostgreSQL stores ISS-0001. Final Search must be case-independent and should return current ACTIVE Issues only."),
}

METHODS = {
    "T41": "BookIssueResponseDto createBookIssue(BookIssueCreateRequestDto request)",
    "T42": "BookIssueResponseDto readBookIssue(Long id)",
    "T43": "BookIssueResponseDto updateBookIssue(Long id, BookIssueUpdateRequestDto request)",
    "T44": "BookIssueResponseDto deleteBookIssue(Long id)",
    "T45": "List<BookIssueResponseDto> searchBookIssue(String text)",
}

NORMAL = {
"T41": '''@Autowired private BookIssueDao bookIssueDao;
@Autowired private BookIssueDtoDoMapper mapper;

@Override
public BookIssueResponseDto createBookIssue(BookIssueCreateRequestDto request) {
    BookIssueDO dataObject = mapper.toDO(request);
    BookIssueDO saved = bookIssueDao.save(dataObject);
    return mapper.toResponse(saved);
}''',
"T42": '''@Autowired private BookIssueDao bookIssueDao;
@Autowired private BookIssueDtoDoMapper mapper;

@Override
public BookIssueResponseDto readBookIssue(Long id) {
    BookIssueDO issue = bookIssueDao.findById(id).orElse(null);
    return mapper.toResponse(issue);
}''',
"T43": '''@Autowired private BookIssueDao bookIssueDao;
@Autowired private BookIssueDtoDoMapper mapper;

@Override
public BookIssueResponseDto updateBookIssue(Long id, BookIssueUpdateRequestDto request) {
    BookIssueDO issue = bookIssueDao.findById(id).orElse(null);
    mapper.applyUpdate(request, issue);
    return mapper.toResponse(bookIssueDao.save(issue));
}''',
"T44": '''@Autowired private BookIssueDao bookIssueDao;
@Autowired private BookIssueDtoDoMapper mapper;

@Override
public BookIssueResponseDto deleteBookIssue(Long id) {
    BookIssueDO issue = bookIssueDao.findById(id).orElse(null);
    issue.setStatus("CANCELLED");
    return mapper.toResponse(bookIssueDao.save(issue));
}''',
"T45": '''@Autowired private BookIssueDao bookIssueDao;
@Autowired private BookIssueDtoDoMapper mapper;

@Override
public List<BookIssueResponseDto> searchBookIssue(String text) {
    String normalized = text.trim();
    return bookIssueDao.search(normalized).stream()
            .map(mapper::toResponse)
            .toList();
}'''
}

OBSERVE = {
"T41": '''@Test
void createWithNullRequestObserveCurrentBehavior() {
    service.createBookIssue(null);
}''',
"T42": '''@Test
void readUnknownBookIssueObserveCurrentBehavior() {
    when(bookIssueDao.findById(999L)).thenReturn(Optional.empty());
    service.readBookIssue(999L);
}''',
"T43": '''@Test
void updateWithNullRequestObserveCurrentBehavior() {
    when(bookIssueDao.findById(1L)).thenReturn(
            Optional.of(issue(1L, "ISS-0001", "ACTIVE")));
    service.updateBookIssue(1L, null);
}''',
"T44": '''@Test
void cancelUnknownBookIssueObserveCurrentBehavior() {
    when(bookIssueDao.findById(999L)).thenReturn(Optional.empty());
    service.deleteBookIssue(999L);
}''',
"T45": '''@Test
void searchWithNullTextObserveCurrentBehavior() {
    service.searchBookIssue(null);
}'''
}

CORRECTED = {
"T41": '''@Override
public BookIssueResponseDto createBookIssue(BookIssueCreateRequestDto request) {
    if (request == null) {
        throw new IllegalArgumentException("Book Issue request is required");
    }
    String issueNumber = request.getIssueNumber().trim();
    bookIssueDao.findByNormalizedBusinessKey(issueNumber)
            .ifPresent(existing -> {
                throw new IllegalStateException("Issue number already exists");
            });
    request.setIssueNumber(issueNumber);
    request.setStatus(request.getStatus().trim());
    return mapper.toResponse(bookIssueDao.save(mapper.toDO(request)));
}''',
"T42": '''@Override
public BookIssueResponseDto readBookIssue(Long id) {
    if (id == null) {
        throw new IllegalArgumentException("Book Issue ID is required");
    }
    BookIssueDO issue = bookIssueDao.findCurrentById(id)
            .orElseThrow(() ->
                    new IllegalArgumentException("Book Issue not found"));
    return mapper.toResponse(issue);
}''',
"T43": '''@Override
public BookIssueResponseDto updateBookIssue(Long id, BookIssueUpdateRequestDto request) {
    if (id == null || request == null) {
        throw new IllegalArgumentException("Book Issue ID and request are required");
    }
    BookIssueDO issue = bookIssueDao.findById(id)
            .orElseThrow(() ->
                    new IllegalArgumentException("Book Issue not found"));
    String issueNumber = request.getIssueNumber().trim();
    bookIssueDao.findByNormalizedBusinessKey(issueNumber)
            .filter(other -> !other.getBookIssueId().equals(id))
            .ifPresent(other -> {
                throw new IllegalStateException("Issue number already exists");
            });
    request.setIssueNumber(issueNumber);
    request.setStatus(request.getStatus().trim());
    mapper.applyUpdate(request, issue);
    return mapper.toResponse(bookIssueDao.save(issue));
}''',
"T44": '''@Override
public BookIssueResponseDto deleteBookIssue(Long id) {
    if (id == null) {
        throw new IllegalArgumentException("Book Issue ID is required");
    }
    BookIssueDO issue = bookIssueDao.findCurrentById(id)
            .orElseThrow(() ->
                    new IllegalArgumentException("Active Book Issue not found"));
    if (bookIssueDao.countDependencies(id) > 0) {
        throw new IllegalStateException("Book Issue already has a Return record");
    }
    issue.setStatus("CANCELLED");
    return mapper.toResponse(bookIssueDao.save(issue));
}''',
"T45": '''@Override
public List<BookIssueResponseDto> searchBookIssue(String text) {
    if (text == null) {
        throw new IllegalArgumentException("Search text is required");
    }
    String normalized = text.trim();
    if (normalized.isBlank()) {
        throw new IllegalArgumentException("Search text is required");
    }
    return bookIssueDao.searchIgnoreCaseCurrent(normalized).stream()
            .map(mapper::toResponse)
            .toList();
}'''
}

INTEGRATION = {
"T41": '''@SpringBootTest
class CreateBookIssueLocalIntegrationTest {
    @Autowired private CreateBookIssueService service;

    @Test
    void shouldRejectPersistedDuplicateIssueNumber() {
        assertThrows(IllegalStateException.class, () ->
            service.createBookIssue(issueRequest(
                "ISS-0001", 4L, 4L, null,
                LocalDate.of(2026, 8, 16),
                LocalDate.of(2026, 8, 30), "ACTIVE")));
    }
}''',
"T42": '''@SpringBootTest
class ReadBookIssueLocalIntegrationTest {
    @Autowired private ReadBookIssueService service;

    @Test
    void shouldReadActiveIssue() {
        BookIssueResponseDto result = service.readBookIssue(1L);
        assertEquals("ISS-0001", result.getIssueNumber());
        assertEquals("ACTIVE", result.getStatus());
    }

    @Test
    void normalReadShouldNotExposeReturnedIssue() {
        assertThrows(IllegalArgumentException.class,
                () -> service.readBookIssue(2L));
    }
}''',
"T43": '''@SpringBootTest
class UpdateBookIssueLocalIntegrationTest {
    @Autowired private UpdateBookIssueService service;

    @Test
    void shouldRejectIssueNumberOwnedByAnotherIssue() {
        assertThrows(IllegalStateException.class, () ->
            service.updateBookIssue(2L, updateRequest(
                "ISS-0001", 2L, 1L, null,
                LocalDate.of(2026, 7, 1),
                LocalDate.of(2026, 7, 15), "RETURNED")));
    }
}''',
"T44": '''@SpringBootTest
class CancelBookIssueLocalIntegrationTest {
    @Autowired private DeleteBookIssueService service;

    @Test
    void shouldCancelActiveIssueWithoutReturn() {
        assertEquals("CANCELLED",
                service.deleteBookIssue(1L).getStatus());
    }

    @Test
    void shouldNotCancelIssueAfterReturnRecordExists() {
        assertThrows(IllegalStateException.class,
                () -> service.deleteBookIssue(5L));
    }
}''',
"T45": '''@SpringBootTest
class SearchBookIssueLocalIntegrationTest {
    @Autowired private SearchBookIssueService service;

    @Test
    void lowercaseSearchFindsUppercaseIssueNumber() {
        assertTrue(service.searchBookIssue("iss-0001").stream()
                .anyMatch(i -> "ISS-0001".equals(i.getIssueNumber())));
    }

    @Test
    void historicalIssueIsNotReturnedByCurrentSearch() {
        assertTrue(service.searchBookIssue("iss-0002").isEmpty());
    }

    @Test
    void validNoMatchReturnsEmptyList() {
        assertTrue(service.searchBookIssue("NO-SUCH-ISSUE").isEmpty());
    }
}'''
}

UPDATED_CODES = {
    "T41": [("00","Book Issue Created Successfully"),("01","Invalid / missing Book Issue request"),("03","Issue number already exists")],
    "T42": [("00","Book Issue Read Successfully"),("01","Book Issue ID is required"),("02","Book Issue not found / historical state not exposed as current")],
    "T43": [("00","Book Issue Updated Successfully"),("01","Invalid Book Issue ID/request"),("02","Book Issue not found"),("03","Issue number belongs to another Book Issue")],
    "T44": [("00","Book Issue Cancelled Successfully"),("01","Book Issue ID is required"),("02","Active Book Issue not found"),("05","Book Issue already has a persisted non-VOID Return")],
    "T45": [("00","Book Issue Search Completed Successfully"),("01","Search text is required or blank")],
}


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def setup(doc):
    for s in doc.sections:
        s.top_margin = Inches(.62)
        s.bottom_margin = Inches(.62)
        s.left_margin = Inches(.72)
        s.right_margin = Inches(.72)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(10.2)
    for n, size in [("Title",21),("Heading 1",15),("Heading 2",12.7),("Heading 3",11.1)]:
        doc.styles[n].font.name = "Aptos"
        doc.styles[n].font.size = Pt(size)


def title(doc, text, subtitle):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.bold = True


def add_table(doc, rows, headers=("Item","Meaning / Value")):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.cell(0,0).text = headers[0]
    t.cell(0,1).text = headers[1]
    shade(t.cell(0,0), "D9EAF7")
    shade(t.cell(0,1), "D9EAF7")
    for a,b in rows:
        c = t.add_row().cells
        c[0].text = str(a)
        c[1].text = str(b)


def add_code(doc, source):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(.12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(textwrap.dedent(source).strip())
    r.font.name = "Consolas"
    r.font.size = Pt(8.05)
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), "F3F4F6")
    p._p.get_or_add_pPr().append(sh)


def note(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    shade(t.cell(0,0), "FFF7D6")
    p = t.cell(0,0).paragraphs[0]
    r = p.add_run("Important: ")
    r.bold = True
    p.add_run(text)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_initial(tid, d):
    doc = Document(); setup(doc)
    title(doc, f"{tid} - {d['operation']} - Initial API Contract",
          f"{d['student']} | {CONTRACT} | Book Issue Module")
    doc.add_heading("1. Fixed Assignment",1)
    add_table(doc, [("Track",tid),("Student",d["student"]),("Operation",d["operation"]),("Service Code",d["code"]),("Endpoint",d["endpoint"]),("Service Interface",d["service"]),("Method",METHODS[tid]),("Success responseCode","00"),("Success Message",d["success"])])
    doc.add_heading("2. JSON Response Envelope",1)
    add_code(doc, f'''{{
  "serviceCode": "{d['code']}",
  "responseCode": "00",
  "message": "{d['success']}",
  "data": {{ }}
}}''')
    doc.add_heading("3. Book Issue Fields",1)
    add_table(doc, [("bookIssueId","Database identifier; response only."),("issueNumber","Business key such as ISS-0001."),("membershipId","Membership that borrows the copy."),("bookCopyId","Physical Book Copy identifier."),("reservationId","Optional Reservation identifier."),("issueDate","Issue date."),("dueDate","Expected return date."),("status","ACTIVE, RETURNED or CANCELLED.")])
    doc.add_heading("4. Training Data",1)
    add_table(doc, [("Issue 1","ISS-0001 / ACTIVE"),("Issue 2","ISS-0002 / RETURNED / Return RET-0001 COMPLETED"),("Issue 3","ISS-0003 / RETURNED / Return RET-0002 VOID"),("Issue 4","ISS-0004 / CANCELLED"),("Issue 5","ISS-0005 / ACTIVE / Return RET-0003 COMPLETED - controlled intermediate workflow state")])
    doc.add_heading("5. Ownership",1)
    add_table(doc, [("Presenter","REST Controller, DTOs, DO, Mapper, DAO base, Flyway, JSON envelope and Thymeleaf framework."),("Student",f"{d['impl']}, track Unit/Integration Tests and {d['frontend']}.")])
    note(doc, "Delete for Book Issue means Cancel. The row remains in PostgreSQL and its status becomes CANCELLED.")
    return doc


def build_progressive(tid, d):
    doc = Document(); setup(doc)
    title(doc, f"{d['student']} - {d['operation']} Progressive Testing Guide",
          f"{tid} | Book Issue Module | {CONTRACT}")
    doc.add_paragraph("Follow the accepted progression: prove the empty service, add the natural successful path, keep its positive Unit Test, execute the next realistic condition before changing the code, understand the observed behavior, correct the implementation, verify coverage, run the real PostgreSQL Integration Test, repeat it with Testcontainers, and finally complete the assigned frontend.")
    doc.add_heading("1. Assignment and Flow",1)
    add_table(doc, [("Backend Assignment",d["operation"]),("Service",d["service"]),("Implementation",d["impl"]),("REST Endpoint",d["endpoint"]),("Service Code",d["code"]),("Frontend Assignment",d["frontend"]),("REST Controller","Presenter-owned")])
    add_table(doc, [("Browser / REST Client","Sends HTTP request."),("Spring Web + Jackson","Creates request DTO for POST/PUT."),("REST Controller","Calls the frozen service."),("Student Service","Business logic."),("Mapper","DTO <-> DO."),("DAO","JPA persistence."),("PostgreSQL","Real Issue/Return state.")])
    note(doc, "The service receives Java DTOs, not raw JSON.")

    doc.add_heading("2. Create the Empty Service",1)
    doc.add_paragraph("The complete Spring chain contains many components. The empty service proves one component independently before HTTP or database dependencies are introduced. Mockito later replaces DAO/Mapper dependencies during Unit Testing.")
    ret = "return java.util.Collections.emptyList();" if tid == "T45" else "return null;"
    add_code(doc, f'''@Service
public class {d['impl']} implements {d['service']} {{
    @Override
    public {METHODS[tid]} {{
        {ret}
    }}
}}''')
    doc.add_heading("2.1 Invoke the empty service",2)
    add_code(doc, f'''@ExtendWith(MockitoExtension.class)
class {d['impl']}Test {{
    @InjectMocks private {d['impl']} service;

    @Test
    void shouldInvokeService() {{
        // Invoke {METHODS[tid]}
        // Only invocation is being proved here.
    }}
}}''')
    bullets(doc, ["Run the focused method from Eclipse using Run As -> JUnit Test.","Do not add the complete business implementation until this test is green."])

    doc.add_heading("3. Add the Natural Successful Implementation",1)
    add_code(doc, NORMAL[tid])
    doc.add_paragraph("Add one focused positive Mockito Unit Test for this successful path. Keep it permanently as a regression test.")

    doc.add_heading("4. Try the Next Realistic Condition",1)
    doc.add_paragraph(d["unit"])
    doc.add_paragraph("Do not modify the service first. Run the test and observe what the current implementation does.")
    add_code(doc, OBSERVE[tid])
    doc.add_heading("4.1 Understand NullPointerException",2)
    doc.add_paragraph("Read the exception type and the first stack-trace line that belongs to the service. A NullPointerException means a reference is null but the code uses it like an object. Correct the application condition before the null value is dereferenced; do not catch NullPointerException everywhere.")

    doc.add_heading("5. Improve the Service",1)
    add_code(doc, CORRECTED[tid])
    doc.add_paragraph("Rerun the newly added test and the earlier positive test. The original success case must remain green.")

    doc.add_heading("6. JaCoCo Coverage",1)
    add_code(doc, "mvnw.cmd clean test")
    add_table(doc, [("Green","Executed by tests."),("Yellow","Partially covered branch."),("Red","Not executed; there is still room for uncertainty.")])
    doc.add_paragraph("Open target/site/jacoco/index.html and inspect the assigned service implementation.")

    doc.add_heading("7. Why Integration Testing Is Required",1)
    doc.add_paragraph("Mockito can simulate a duplicate issue number, missing row or Return dependency. It cannot prove the real Flyway data, JPA mapping, lifecycle states, PostgreSQL string comparison or Issue-to-Return relationship. Integration Testing joins those real components.")
    add_table(doc, [("Unit Test","Service + Mockito mocks."),("Local Integration","Spring + real Mapper + DAO + local PostgreSQL."),("Testcontainers","Same path with temporary PostgreSQL 18.")])

    doc.add_heading("8. PostgreSQL JDBC Configuration",1)
    add_code(doc, '''spring.datasource.url=jdbc:postgresql://localhost:5432/library_full_stack
spring.datasource.username=postgres
spring.datasource.password=postgres
spring.datasource.driver-class-name=org.postgresql.Driver''')
    add_code(doc, '''<dependency>
    <groupId>org.postgresql</groupId>
    <artifactId>postgresql</artifactId>
    <scope>runtime</scope>
</dependency>

<!-- MySQL learning reference only - keep commented
<dependency>
    <groupId>com.mysql</groupId>
    <artifactId>mysql-connector-j</artifactId>
    <scope>runtime</scope>
</dependency>
-->''')
    add_code(doc, '''# MySQL reference only - DO NOT activate
# spring.datasource.url=jdbc:mysql://localhost:3306/library_full_stack
# spring.datasource.driver-class-name=com.mysql.cj.jdbc.Driver''')

    doc.add_heading("9. Inspect Real Data in DBeaver",1)
    add_code(doc, '''SELECT book_issue_id, issue_number, fk_membership, fk_book_copy,
       fk_reservation, issue_date, due_date, status
FROM tbl_book_issue
ORDER BY book_issue_id;

SELECT book_return_id, return_number, fk_book_issue, return_date, status
FROM tbl_book_return
ORDER BY book_return_id;''')
    doc.add_paragraph(d["integration"])
    if tid == "T44":
        note(doc, "ISS-0005 is a controlled intermediate workflow state: the Return is already persisted while the Issue status update is still pending. The dependency guard prevents cancellation after a Return exists.")

    doc.add_heading("10. Complete Local PostgreSQL Integration Test",1)
    add_code(doc, INTEGRATION[tid])
    if tid == "T45":
        doc.add_heading("10.1 Corrected Search Query",2)
        add_code(doc, '''@Transactional(readOnly = true)
public List<BookIssueDO> searchIgnoreCaseCurrent(String text) {
    String pattern = "%" + text.toLowerCase() + "%";
    return entityManager.createQuery(
            "select i from BookIssueDO i " +
            "where i.status = 'ACTIVE' and lower(i.issueNumber) like :pattern " +
            "order by i.bookIssueId", BookIssueDO.class)
            .setParameter("pattern", pattern)
            .getResultList();
}''')

    doc.add_heading("11. Complete PostgreSQL Testcontainers Test",1)
    tc = INTEGRATION[tid].replace("@SpringBootTest", "@SpringBootTest\n@Testcontainers")
    tc = tc.replace("{", '''{
    @Container
    @ServiceConnection
    static PostgreSQLContainer postgres =
        new PostgreSQLContainer(
            DockerImageName.parse("postgres:18"));
''', 1)
    add_code(doc, tc)
    doc.add_paragraph("Flyway creates the deterministic training rows inside the temporary container, so the same persistence behavior is tested on another machine.")

    doc.add_heading("12. Final Regression",1)
    add_code(doc, "mvnw.cmd clean test")
    bullets(doc, ["Focused JUnit test green.","Original successful Unit Test still green.","Local PostgreSQL Integration Test green.","PostgreSQL Testcontainers test green.","JaCoCo reviewed."])

    doc.add_heading("13. Student-owned Files",1)
    add_table(doc, [("Service",d["impl"]+".java"),("Tests","Track-specific Unit and Integration Test classes."),("Frontend",d["frontend"]),("Presenter-owned","REST Controller, DTOs, DO, Mapper, base DAO, Flyway and Thymeleaf framework.")])
    if tid == "T45":
        note(doc, "T45 owns only the case-independent current-Issue search DAO method named in this guide.")

    doc.add_heading("14. Frontend Assignment",1)
    doc.add_paragraph("Build a working page first in one JavaScript file. Prove main.js execution and the button handler before the REST call. Test against Presenter STUB/support, then refactor into main.js, API, form and view modules. The partner integration must keep the same /rest URL.")
    add_code(doc, '''<script type="module" src="/src/main.js"></script>

console.log("main.js loaded");
document.querySelector("#actionButton")
    .addEventListener("click", () => {
        console.log("action button reached JavaScript");
    });''')
    doc.add_heading("14.1 Transport handling",2)
    add_code(doc, '''const response = await fetch(url, options);
if (response.status === 404) {
  showMessage("NOT Found");
  return;
}
if (!response.ok) {
  showMessage("Error Encountered, Please contact Administrator");
  return;
}
const body = await response.json();
showMessage(body.message);''')
    add_table(doc, [("src/main.js","Starts the page."),("src/api/bookIssueApi.js","Contains REST calls."),("src/forms/bookIssueForm.js","Reads input."),("src/views/bookIssueView.js","Renders messages and rows.")])
    bullets(doc, ["Presenter STUB/support works.","Page is refactored only after it works.","Partner service works with unchanged endpoint.","Assigned files are committed and pushed."])
    return doc


def build_updated(tid, d):
    doc = Document(); setup(doc)
    title(doc, f"{tid} - {d['operation']} - Updated API Contract",
          f"{CONTRACT} | Book Issue Module")
    doc.add_heading("1. Fixed Endpoint",1)
    add_table(doc, [("Track",tid),("Service Code",d["code"]),("Endpoint",d["endpoint"]),("Controlled application result","HTTP 200")])
    doc.add_heading("2. Frozen JSON Envelope",1)
    add_code(doc, '''{
  "serviceCode": "XX",
  "responseCode": "00",
  "message": "...",
  "data": {}
}''')
    doc.add_heading("3. Application Response Codes",1)
    add_table(doc, UPDATED_CODES[tid], ("responseCode","Meaning"))
    doc.add_heading("4. Frontend Transport Errors",1)
    add_table(doc, [("HTTP 404","NOT Found"),("Other non-200","Error Encountered, Please contact Administrator")])
    note(doc, "A valid Search with zero matches returns responseCode 00 and an empty data list." if tid == "T45" else "Controlled application results use the frozen JSON envelope with HTTP 200.")
    return doc


OUT.mkdir(parents=True, exist_ok=True)
for tid, data in TRACKS.items():
    folder = OUT / f"{tid}_{data['operation'].replace(' ','_')}"
    folder.mkdir(parents=True, exist_ok=True)
    build_initial(tid, data).save(folder / f"{tid}_01_Initial_API_Contract.docx")
    build_progressive(tid, data).save(folder / f"{tid}_02_Progressive_Development_Guide.docx")
    build_updated(tid, data).save(folder / f"{tid}_03_Updated_API_Contract.docx")

print("Generated T41-T45 Book Issue student guides")
