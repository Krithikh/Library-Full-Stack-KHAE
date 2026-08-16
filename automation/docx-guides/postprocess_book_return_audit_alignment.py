from pathlib import Path

from docx import Document


ROOT = Path("Documents/Student-Guides/Book-Return")


def replace_in_paragraph(paragraph, replacements):
    original = paragraph.text
    updated = original
    for old, new in replacements.items():
        updated = updated.replace(old, new)

    if updated == original:
        return False

    # These generated guide paragraphs use simple formatting.  Replacing the
    # paragraph text is intentionally limited to the audit-alignment strings.
    paragraph.text = updated
    return True


def replace_everywhere(doc, replacements):
    changed = False

    for paragraph in doc.paragraphs:
        changed = replace_in_paragraph(paragraph, replacements) or changed

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changed = replace_in_paragraph(paragraph, replacements) or changed

    return changed


def align_initial_training_table(path):
    doc = Document(path)
    changed = False

    for table in doc.tables:
        has_ret_0003 = False
        has_ret_0004 = False

        for row in table.rows:
            if len(row.cells) < 2:
                continue

            key = row.cells[0].text.strip()
            if key == "RET-0003":
                has_ret_0003 = True
                expected = "COMPLETED; FINE-0003 OUTSTANDING"
                if row.cells[1].text.strip() != expected:
                    row.cells[1].text = expected
                    changed = True
            elif key == "RET-0004":
                has_ret_0004 = True

        if has_ret_0003 and not has_ret_0004:
            row = table.add_row().cells
            row[0].text = "RET-0004"
            row[1].text = "COMPLETED; no Fine - T49 successful Void case"
            changed = True

    if changed:
        doc.save(path)


def align_t49_progressive(path):
    doc = Document(path)

    replacements = {
        "RET-0001 has non-VOID Fine FINE-0001; RET-0003 has no Fine":
            "RET-0001 has non-VOID Fine FINE-0001; RET-0004 has no Fine",
        "service.deleteBookReturn(3L)":
            "service.deleteBookReturn(4L)",
        "RET-0001, RET-0002, RET-0003 and Fine relationships":
            "RET-0001, RET-0002, RET-0003, RET-0004 and Fine relationships",
    }

    if replace_everywhere(doc, replacements):
        doc.save(path)


def main():
    for path in ROOT.glob("T4*_*/T4*_01_Initial_API_Contract.docx"):
        align_initial_training_table(path)

    align_t49_progressive(
        ROOT
        / "T49_Void_Book_Return"
        / "T49_02_Progressive_Development_Guide.docx"
    )

    print("Applied final T49 audit alignment to generated Book Return guides.")


if __name__ == "__main__":
    main()
