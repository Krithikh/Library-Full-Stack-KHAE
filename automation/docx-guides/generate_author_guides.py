from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import textwrap

OUT = Path('Documents/Student-Guides/Author')
CONTRACT = 'LIBRARY-FULLSTACK-V1'

TRACKS = {
 'T11': dict(student='Student 11', operation='Create Author', code='11', endpoint='POST /rest/authors', service='CreateAuthorService', impl='CreateAuthorServiceImpl', success='Author Created Successfully', frontend='Update Author UI', unit='Call createAuthor(null) after the successful Create path is working.', integration='The database already contains AUT-001. Try to create another Author using AUT-001.'),
 'T12': dict(student='Student 12', operation='Read Author', code='12', endpoint='GET /rest/authors/{id}', service='ReadAuthorService', impl='ReadAuthorServiceImpl', success='Author Read Successfully', frontend='Deactivate Author UI', unit='Request Author ID 999 after the successful Read path is working.', integration='Author ID 3 physically exists as AUT-OLD but is inactive. Normal Read must not expose it.'),
 'T13': dict(student='Student 13', operation='Update Author', code='13', endpoint='PUT /rest/authors/{id}', service='UpdateAuthorService', impl='UpdateAuthorServiceImpl', success='Author Updated Successfully', frontend='Search Author UI', unit='Call updateAuthor(2L, null) after the successful Update path is working.', integration='Author ID 1 owns AUT-001. Try to update Author ID 2 from AUT-002 to AUT-001.'),
 'T14': dict(student='Student 14', operation='Deactivate Author', code='14', endpoint='DELETE /rest/authors/{id}', service='DeleteAuthorService', impl='DeleteAuthorServiceImpl', success='Author Deactivated Successfully', frontend='Create Author UI', unit='Request Author ID 999 after logical deactivation works for a valid Author.', integration='Author ID 1 is referenced by active Book 1. It must not be deactivated while the active Book dependency exists.'),
 'T15': dict(student='Student 15', operation='Search Author', code='15', endpoint='GET /rest/authors/search?text={text}', service='SearchAuthorService', impl='SearchAuthorServiceImpl', success='Author Search Completed Successfully', frontend='Author List / Search Results UI', unit='Call searchAuthor(null) after the successful Search path is working.', integration='Search for aut-001 and robert while PostgreSQL stores AUT-001 and Robert C. Martin.')
}

METHODS = {
 'T11':'AuthorResponseDto createAuthor(AuthorCreateRequestDto request)',
 'T12':'AuthorResponseDto readAuthor(Long id)',
 'T13':'AuthorResponseDto updateAuthor(Long id, AuthorUpdateRequestDto request)',
 'T14':'AuthorResponseDto deleteAuthor(Long id)',
 'T15':'List<AuthorResponseDto> searchAuthor(String text)'
}

NORMAL = {
 'T11':'''@Autowired private AuthorDao authorDao;\n@Autowired private AuthorDtoDoMapper mapper;\n\n@Override\npublic AuthorResponseDto createAuthor(AuthorCreateRequestDto request) {\n    AuthorDO dataObject = mapper.toDO(request);\n    AuthorDO saved = authorDao.save(dataObject);\n    return mapper.toResponse(saved);\n}''',
 'T12':'''@Autowired private AuthorDao authorDao;\n@Autowired private AuthorDtoDoMapper mapper;\n\n@Override\npublic AuthorResponseDto readAuthor(Long id) {\n    AuthorDO author = authorDao.findById(id).orElse(null);\n    return mapper.toResponse(author);\n}''',
 'T13':'''@Autowired private AuthorDao authorDao;\n@Autowired private AuthorDtoDoMapper mapper;\n\n@Override\npublic AuthorResponseDto updateAuthor(Long id, AuthorUpdateRequestDto request) {\n    AuthorDO author = authorDao.findById(id).orElse(null);\n    mapper.applyUpdate(request, author);\n    return mapper.toResponse(authorDao.save(author));\n}''',
 'T14':'''@Autowired private AuthorDao authorDao;\n@Autowired private AuthorDtoDoMapper mapper;\n\n@Override\npublic AuthorResponseDto deleteAuthor(Long id) {\n    AuthorDO author = authorDao.findById(id).orElse(null);\n    author.setActive(false);\n    return mapper.toResponse(authorDao.save(author));\n}''',
 'T15':'''@Autowired private AuthorDao authorDao;\n@Autowired private AuthorDtoDoMapper mapper;\n\n@Override\npublic List<AuthorResponseDto> searchAuthor(String text) {\n    String normalized = text.trim();\n    return authorDao.search(normalized).stream()\n            .map(mapper::toResponse)\n            .toList();\n}'''
}

OBSERVE = {
 'T11':'''@Test\nvoid createWithNullRequestObserveCurrentBehavior() {\n    service.createAuthor(null);\n}''',
 'T12':'''@Test\nvoid readUnknownAuthorObserveCurrentBehavior() {\n    when(authorDao.findById(999L)).thenReturn(Optional.empty());\n    service.readAuthor(999L);\n}''',
 'T13':'''@Test\nvoid updateWithNullRequestObserveCurrentBehavior() {\n    when(authorDao.findById(2L)).thenReturn(Optional.of(author(2L, "AUT-002", true)));\n    service.updateAuthor(2L, null);\n}''',
 'T14':'''@Test\nvoid deactivateUnknownAuthorObserveCurrentBehavior() {\n    when(authorDao.findById(999L)).thenReturn(Optional.empty());\n    service.deleteAuthor(999L);\n}''',
 'T15':'''@Test\nvoid searchWithNullTextObserveCurrentBehavior() {\n    service.searchAuthor(null);\n}'''
}

CORRECTED = {
 'T11':'''@Override\npublic AuthorResponseDto createAuthor(AuthorCreateRequestDto request) {\n    if (request == null) {\n        throw new IllegalArgumentException("Author request is required");\n    }\n    String code = request.getAuthorCode().trim();\n    authorDao.findByNormalizedBusinessKey(code).ifPresent(existing -> {\n        throw new IllegalStateException("Author code already exists");\n    });\n    request.setAuthorCode(code);\n    request.setAuthorName(request.getAuthorName().trim());\n    return mapper.toResponse(authorDao.save(mapper.toDO(request)));\n}''',
 'T12':'''@Override\npublic AuthorResponseDto readAuthor(Long id) {\n    if (id == null) {\n        throw new IllegalArgumentException("Author ID is required");\n    }\n    AuthorDO author = authorDao.findCurrentById(id)\n            .orElseThrow(() -> new IllegalArgumentException("Author not found"));\n    return mapper.toResponse(author);\n}''',
 'T13':'''@Override\npublic AuthorResponseDto updateAuthor(Long id, AuthorUpdateRequestDto request) {\n    if (id == null || request == null) {\n        throw new IllegalArgumentException("Author ID and request are required");\n    }\n    AuthorDO author = authorDao.findCurrentById(id)\n            .orElseThrow(() -> new IllegalArgumentException("Author not found"));\n    String code = request.getAuthorCode().trim();\n    authorDao.findByNormalizedBusinessKey(code)\n            .filter(other -> !other.getAuthorId().equals(id))\n            .ifPresent(other -> {\n                throw new IllegalStateException("Author code already exists");\n            });\n    request.setAuthorCode(code);\n    request.setAuthorName(request.getAuthorName().trim());\n    mapper.applyUpdate(request, author);\n    return mapper.toResponse(authorDao.save(author));\n}''',
 'T14':'''@Override\npublic AuthorResponseDto deleteAuthor(Long id) {\n    if (id == null) {\n        throw new IllegalArgumentException("Author ID is required");\n    }\n    AuthorDO author = authorDao.findCurrentById(id)\n            .orElseThrow(() -> new IllegalArgumentException("Author not found"));\n    if (authorDao.countDependencies(id) > 0) {\n        throw new IllegalStateException("Author is used by an active Book");\n    }\n    author.setActive(false);\n    return mapper.toResponse(authorDao.save(author));\n}''',
 'T15':'''@Override\npublic List<AuthorResponseDto> searchAuthor(String text) {\n    if (text == null) {\n        throw new IllegalArgumentException("Search text is required");\n    }\n    String normalized = text.trim();\n    if (normalized.isBlank()) {\n        throw new IllegalArgumentException("Search text is required");\n    }\n    return authorDao.searchIgnoreCase(normalized).stream()\n            .map(mapper::toResponse)\n            .toList();\n}'''
}

INTEGRATION = {
 'T11':'''@SpringBootTest\nclass CreateAuthorLocalIntegrationTest {\n    @Autowired private CreateAuthorService service;\n\n    @Test\n    void shouldCreateNewAuthorCode() {\n        AuthorResponseDto result = service.createAuthor(request("AUT-NEW", "New Author"));\n        assertEquals("AUT-NEW", result.getAuthorCode());\n    }\n\n    @Test\n    void shouldRejectPersistedDuplicateAuthorCode() {\n        assertThrows(IllegalStateException.class,\n                () -> service.createAuthor(request("AUT-001", "Another Author")));\n    }\n}''',
 'T12':'''@SpringBootTest\nclass ReadAuthorLocalIntegrationTest {\n    @Autowired private ReadAuthorService service;\n\n    @Test\n    void shouldReadActiveAuthor() {\n        assertEquals("AUT-002", service.readAuthor(2L).getAuthorCode());\n    }\n\n    @Test\n    void normalReadShouldNotExposeInactiveAuthor() {\n        assertThrows(IllegalArgumentException.class, () -> service.readAuthor(3L));\n    }\n}''',
 'T13':'''@SpringBootTest\nclass UpdateAuthorLocalIntegrationTest {\n    @Autowired private UpdateAuthorService service;\n\n    @Test\n    void shouldAllowCurrentAuthorToKeepOwnCode() {\n        AuthorResponseDto result = service.updateAuthor(1L, request("AUT-001", "Robert C. Martin"));\n        assertEquals("AUT-001", result.getAuthorCode());\n    }\n\n    @Test\n    void shouldRejectCodeOwnedByAnotherAuthor() {\n        assertThrows(IllegalStateException.class,\n                () -> service.updateAuthor(2L, request("AUT-001", "Joshua Bloch")));\n    }\n}''',
 'T14':'''@SpringBootTest\nclass DeleteAuthorLocalIntegrationTest {\n    @Autowired private DeleteAuthorService service;\n\n    @Test\n    void shouldDeactivateUnusedAuthor() {\n        assertFalse(service.deleteAuthor(4L).getActive());\n    }\n\n    @Test\n    void shouldNotDeactivateAuthorUsedByActiveBook() {\n        assertThrows(IllegalStateException.class, () -> service.deleteAuthor(1L));\n    }\n}''',
 'T15':'''@SpringBootTest\nclass SearchAuthorLocalIntegrationTest {\n    @Autowired private SearchAuthorService service;\n\n    @Test\n    void shouldFindStoredUppercaseCodeUsingLowercaseSearch() {\n        assertTrue(service.searchAuthor("aut-001").stream()\n                .anyMatch(a -> "AUT-001".equals(a.getAuthorCode())));\n    }\n\n    @Test\n    void shouldFindMixedCaseNameUsingLowercaseSearch() {\n        assertTrue(service.searchAuthor("robert").stream()\n                .anyMatch(a -> "Robert C. Martin".equals(a.getAuthorName())));\n    }\n}'''
}

UPDATED = {
 'T11':[('00','Author Created Successfully'),('01','Invalid / missing Author request'),('03','Author code already exists')],
 'T12':[('00','Author Read Successfully'),('01','Author ID is required'),('02','Author not found / inactive row is not exposed')],
 'T13':[('00','Author Updated Successfully'),('01','Invalid Author ID/request'),('02','Author not found'),('03','Author code belongs to another Author')],
 'T14':[('00','Author Deactivated Successfully'),('01','Author ID is required'),('02','Author not found'),('05','Author is used by an active Book')],
 'T15':[('00','Author Search Completed Successfully'),('01','Search text is required or blank')]
}


def shade(cell, fill):
    p = cell._tc.get_or_add_tcPr(); s = OxmlElement('w:shd'); s.set(qn('w:fill'), fill); p.append(s)


def setup(doc):
    for section in doc.sections:
        section.top_margin = Inches(.65); section.bottom_margin = Inches(.65); section.left_margin = Inches(.72); section.right_margin = Inches(.72)
    doc.styles['Normal'].font.name = 'Aptos'; doc.styles['Normal'].font.size = Pt(10.2)
    for n, sz in [('Title',21),('Heading 1',15),('Heading 2',12.7),('Heading 3',11.2)]:
        doc.styles[n].font.name = 'Aptos'; doc.styles[n].font.size = Pt(sz)


def add_title(doc, a, b):
    p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(a)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(b); r.bold=True


def add_table(doc, rows, headers=('Item','Meaning / Value')):
    t=doc.add_table(rows=1, cols=2); t.style='Table Grid'; t.cell(0,0).text=headers[0]; t.cell(0,1).text=headers[1]; shade(t.cell(0,0),'D9EAF7'); shade(t.cell(0,1),'D9EAF7')
    for a,b in rows:
        c=t.add_row().cells; c[0].text=str(a); c[1].text=str(b)


def add_code(doc, source):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.13); r=p.add_run(textwrap.dedent(source).strip()); r.font.name='Consolas'; r.font.size=Pt(8.2); s=OxmlElement('w:shd'); s.set(qn('w:fill'),'F3F4F6'); p._p.get_or_add_pPr().append(s)


def add_note(doc, text):
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'; shade(t.cell(0,0),'FFF7D6'); t.cell(0,0).text=text


def add_bullets(doc, items):
    for x in items: doc.add_paragraph(x, style='List Bullet')


def build_initial(t,d):
    doc=Document(); setup(doc); add_title(doc,f"{t} - {d['operation']} - Initial API Contract",f"{d['student']} | {CONTRACT}")
    doc.add_heading('1. Assignment',1); add_table(doc,[('Track',t),('Student',d['student']),('Operation',d['operation']),('Service Code',d['code']),('Endpoint',d['endpoint']),('Service Interface',d['service']),('Method',METHODS[t]),('Success Response Code','00'),('Success Message',d['success'])])
    doc.add_heading('2. Frozen JSON Response Envelope',1); add_code(doc,f'''{{\n  "serviceCode": "{d['code']}",\n  "responseCode": "00",\n  "message": "{d['success']}",\n  "data": {{ }}\n}}''')
    doc.add_heading('3. Author Fields',1); add_table(doc,[('authorCode','Business key, e.g. AUT-001'),('authorName','Author display name'),('authorId','Database identifier in response'),('active','Logical active state in response')])
    doc.add_heading('4. Training Data',1); add_table(doc,[('1','AUT-001 / Robert C. Martin / ACTIVE'),('2','AUT-002 / Joshua Bloch / ACTIVE'),('3','AUT-OLD / Archived Author / INACTIVE'),('4','AUT-004 / Unused Author / ACTIVE')])
    add_note(doc,'The REST Controller is Presenter-owned. The student implements the assigned service and tests without renaming the endpoint or service code.')
    return doc


def build_progressive(t,d):
    doc=Document(); setup(doc); add_title(doc,f"{d['student']} - {d['operation']} Progressive Testing Guide",f"{t} | Author Module | {CONTRACT}")
    doc.add_paragraph('Follow the sequence in order: prove invocation, implement the normal successful path, test it, try the next realistic condition before changing the service, understand the current behavior, improve the service, and then verify database-dependent behavior with PostgreSQL.')
    doc.add_heading('1. Assignment and Application Flow',1); add_table(doc,[('Track',t),('Backend',d['operation']),('Service',d['service']),('Method',METHODS[t]),('REST Endpoint',d['endpoint']),('Service Code',d['code']),('Frontend',d['frontend']),('REST Controller','Presenter-owned')])
    add_table(doc,[('Browser / REST client','HTTP request'),('Spring Web + Jackson','JSON to DTO'),('REST Controller','Calls fixed service interface'),('Service','Student business behavior'),('AuthorDtoDoMapper','DTO to DO / DO to DTO'),('AuthorDao','JPA persistence'),('PostgreSQL','Real database')])
    doc.add_heading('2. Create the Empty Service',1); doc.add_paragraph('The empty service proves that one component can be constructed and invoked before the complete Spring chain is used. Mockito replaces DAO and Mapper dependencies for Unit Testing; real dependencies are used during Integration Testing.')
    ret='return java.util.Collections.emptyList();' if t=='T15' else 'return null;'; add_code(doc,f'''@Service\npublic class {d['impl']} implements {d['service']} {{\n    @Override\n    public {METHODS[t]} {{\n        {ret}\n    }}\n}}''')
    doc.add_heading('3. Invoke the Empty Service with JUnit',1); add_code(doc,f'''@ExtendWith(MockitoExtension.class)\nclass {d['impl']}Test {{\n    @InjectMocks\n    private {d['impl']} service;\n\n    @Test\n    void shouldInvokeService() {{\n        // Invoke the assigned method here.\n    }}\n}}'''); add_bullets(doc,['Right-click the test in Eclipse.','Choose Run As -> JUnit Test.','Confirm the method is reached before adding business logic.'])
    doc.add_heading(f"4. Implement the Normal {d['operation']} Behavior",1); add_code(doc,NORMAL[t])
    doc.add_heading('5. Positive Unit Test',1); doc.add_paragraph('Create a Mockito test for the successful path. Configure DAO/Mapper mocks, call the service, assert the returned Author values, and verify the expected DAO call. Keep this test after every later correction.')
    doc.add_heading('6. Try the Next Realistic Condition',1); doc.add_paragraph(d['unit']); add_code(doc,OBSERVE[t]); doc.add_paragraph('Run this test before changing the service. Read the exception type and the first service line in the stack trace. A NullPointerException means a null reference was used as though an object existed; handle the application condition before dereferencing it.')
    doc.add_heading('7. Improve the Service',1); add_code(doc,CORRECTED[t]); doc.add_paragraph('Rerun the original successful test and the new test after the correction.')
    doc.add_heading('8. JaCoCo Coverage',1); add_code(doc,'mvnw.cmd clean test'); add_table(doc,[('Green','Executed'),('Yellow','Partially covered'),('Red','Not executed; room for uncertainty')]); doc.add_paragraph('Open target/site/jacoco/index.html and inspect the Author service implementation.')
    doc.add_heading('9. Integration Testing',1); add_table(doc,[('Unit Test','Service + Mockito mocks'),('Local Integration','Spring + real Mapper + DAO + local PostgreSQL'),('Testcontainers','Spring + real Mapper + DAO + temporary PostgreSQL')]); doc.add_paragraph(d['integration'])
    doc.add_heading('10. PostgreSQL JDBC Driver',1); add_code(doc,'''spring.datasource.url=jdbc:postgresql://localhost:5432/library_full_stack\nspring.datasource.username=postgres\nspring.datasource.password=postgres\nspring.datasource.driver-class-name=org.postgresql.Driver'''); add_code(doc,'''<dependency>\n    <groupId>org.postgresql</groupId>\n    <artifactId>postgresql</artifactId>\n    <scope>runtime</scope>\n</dependency>\n\n<!-- MySQL learning reference only - keep commented\n<dependency>\n    <groupId>com.mysql</groupId>\n    <artifactId>mysql-connector-j</artifactId>\n    <scope>runtime</scope>\n</dependency>\n-->'''); add_code(doc,'''# MySQL reference only - DO NOT activate\n# spring.datasource.url=jdbc:mysql://localhost:3306/library_full_stack\n# spring.datasource.driver-class-name=com.mysql.cj.jdbc.Driver''')
    doc.add_heading('11. Inspect Database in DBeaver',1); add_code(doc,'''SELECT author_id, author_code, author_name, is_active\nFROM tbl_author\nORDER BY author_id;\n\nSELECT book_id, isbn, title, fk_author, is_active\nFROM tbl_book\nORDER BY book_id;''')
    doc.add_heading('12. Complete Local Integration Test',1); add_code(doc,INTEGRATION[t])
    if t=='T15':
        doc.add_heading('12.1 Corrected Case-Independent Search',2); add_code(doc,'''@Transactional(readOnly = true)\npublic List<AuthorDO> searchIgnoreCase(String text) {\n    String pattern = "%" + text.toLowerCase() + "%";\n    return entityManager.createQuery(\n            "select a from AuthorDO a where a.active = true and " +\n            "(lower(a.authorCode) like :pattern or lower(a.authorName) like :pattern) " +\n            "order by a.authorId", AuthorDO.class)\n            .setParameter("pattern", pattern)\n            .getResultList();\n}''')
    doc.add_heading('13. PostgreSQL Testcontainers',1); add_code(doc,'''@Container\n@ServiceConnection\nstatic PostgreSQLContainer postgres =\n        new PostgreSQLContainer(DockerImageName.parse("postgres:18"));'''); doc.add_paragraph('Add @Testcontainers to the integration-test class and keep the same integration-test cases. Flyway recreates the known data inside the temporary PostgreSQL database.')
    doc.add_heading('14. Frontend Assignment',1); doc.add_paragraph(f"{d['student']} implements {d['frontend']}. Prove main.js and the button handler first, build a working single-file page, then refactor into main.js, api/authorApi.js, form and view modules."); add_code(doc,'''<script type="module" src="/src/main.js"></script>\n\nconsole.log("main.js loaded");\ndocument.querySelector("#actionButton").addEventListener("click", () => {\n  console.log("action button reached JavaScript");\n});'''); add_code(doc,'''if (response.status === 404) {\n  showMessage("NOT Found");\n  return;\n}\nif (!response.ok) {\n  showMessage("Error Encountered, Please contact Administrator");\n  return;\n}\nconst body = await response.json();\nshowMessage(body.message);'''); doc.add_paragraph('Test first with Presenter STUB/support. Later integrate the partner service without changing the /rest URL.')
    doc.add_heading('15. Completion Checklist',1); add_bullets(doc,['Empty service invoked','Positive Unit Test retained','Next realistic condition observed before correction','Service improved','JaCoCo reviewed','Local PostgreSQL test passed','PostgreSQL Testcontainers test passed','Frontend works with Presenter STUB/support','Partner integration uses the same URL'])
    return doc


def build_updated(t,d):
    doc=Document(); setup(doc); add_title(doc,f"{t} - {d['operation']} - Updated API Contract",f"{CONTRACT} | Released after progressive development"); doc.add_heading('1. Endpoint',1); add_table(doc,[('Service Code',d['code']),('Endpoint',d['endpoint']),('Controlled application result','HTTP 200')]); doc.add_heading('2. Response Codes',1); add_table(doc,UPDATED[t],('responseCode','Meaning')); doc.add_heading('3. Transport Errors',1); add_table(doc,[('HTTP 404','NOT Found'),('Other non-200','Error Encountered, Please contact Administrator')]); add_note(doc,'A valid Search with zero matches returns responseCode 00 with an empty data list.'); return doc


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    for t,d in TRACKS.items():
        folder=OUT/f"{t}_{d['operation'].replace(' ','_')}"; folder.mkdir(parents=True,exist_ok=True)
        build_initial(t,d).save(folder/f'{t}_01_Initial_API_Contract.docx')
        build_progressive(t,d).save(folder/f'{t}_02_Progressive_Development_Guide.docx')
        build_updated(t,d).save(folder/f'{t}_03_Updated_API_Contract.docx')
    print('Generated 15 Author DOCX guides in', OUT)

if __name__ == '__main__': main()
