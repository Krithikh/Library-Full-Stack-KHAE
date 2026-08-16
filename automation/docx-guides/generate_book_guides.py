from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('Documents/Student-Guides/Book')
OUT.mkdir(parents=True, exist_ok=True)
CONTRACT = 'LIBRARY-FULLSTACK-V1'

TRACKS = {
'T26': dict(student='Student 26', op='Create Book', code='26', service='CreateBookService', impl='CreateBookServiceImpl', endpoint='POST /rest/books', success='Book Created Successfully', ui='Update Book UI', unit='Call createBook(null) after the successful Create test works.', integration='The database already contains ISBN 9780132350884 for Clean Code. Try to create another Book with the same persisted ISBN.'),
'T27': dict(student='Student 27', op='Read Book', code='27', service='ReadBookService', impl='ReadBookServiceImpl', endpoint='GET /rest/books/{id}', success='Book Read Successfully', ui='Deactivate Book UI', unit='Request Book ID 999 after the successful Read test works.', integration='Book ID 3 physically exists as Archived Programming Book but is inactive. Normal Read must not expose it.'),
'T28': dict(student='Student 28', op='Update Book', code='28', service='UpdateBookService', impl='UpdateBookServiceImpl', endpoint='PUT /rest/books/{id}', success='Book Updated Successfully', ui='Search Book UI', unit='Call updateBook(2L, null) after the successful Update test works.', integration='Book ID 1 owns ISBN 9780132350884. Try to update Book ID 2 to the same ISBN. Exclude the current row from duplicate detection.'),
'T29': dict(student='Student 29', op='Deactivate Book', code='29', service='DeleteBookService', impl='DeleteBookServiceImpl', endpoint='DELETE /rest/books/{id}', success='Book Deactivated Successfully', ui='Create Book UI', unit='Request Book ID 999 after normal logical deactivation works.', integration='Book ID 2 has only a WITHDRAWN copy and can be deactivated. Book ID 1 has non-withdrawn Book Copies and must not be deactivated.'),
'T30': dict(student='Student 30', op='Search Book', code='30', service='SearchBookService', impl='SearchBookServiceImpl', endpoint='GET /rest/books/search?text={text}', success='Book Search Completed Successfully', ui='Book List / Search Results UI', unit='Call searchBook(null) after normal Search works. The first implementation calls trim().', integration='Search for clean code while PostgreSQL stores Clean Code. The completed search must be active-only and case-independent.'),
}

def shade(cell, fill):
    pr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),fill); pr.append(sh)

def setup(doc):
    for s in doc.sections:
        s.top_margin=Inches(.62); s.bottom_margin=Inches(.62); s.left_margin=Inches(.72); s.right_margin=Inches(.72)
    doc.styles['Normal'].font.name='Aptos'; doc.styles['Normal'].font.size=Pt(10.2)
    for n,sz in [('Title',21),('Heading 1',15),('Heading 2',12.7),('Heading 3',11.1)]:
        doc.styles[n].font.name='Aptos'; doc.styles[n].font.size=Pt(sz)

def title(doc,a,b):
    p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(a)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(b); r.bold=True

def table(doc, rows, headers=('Item','Meaning / Value')):
    t=doc.add_table(rows=1,cols=2); t.style='Table Grid'; t.cell(0,0).text=headers[0]; t.cell(0,1).text=headers[1]; shade(t.cell(0,0),'D9EAF7'); shade(t.cell(0,1),'D9EAF7')
    for a,b in rows:
        c=t.add_row().cells; c[0].text=str(a); c[1].text=str(b)

def code(doc,s):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.12); r=p.add_run(s.strip()); r.font.name='Consolas'; r.font.size=Pt(8.05); sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),'F3F4F6'); p._p.get_or_add_pPr().append(sh)

def note(doc,s):
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'; shade(t.cell(0,0),'FFF7D6'); p=t.cell(0,0).paragraphs[0]; r=p.add_run('Important: '); r.bold=True; p.add_run(s)

def bullets(doc, items):
    for x in items: doc.add_paragraph(x,style='List Bullet')

def sig(t):
    return {'T26':'BookResponseDto createBook(BookCreateRequestDto request)','T27':'BookResponseDto readBook(Long id)','T28':'BookResponseDto updateBook(Long id, BookUpdateRequestDto request)','T29':'BookResponseDto deleteBook(Long id)','T30':'List<BookResponseDto> searchBook(String text)'}[t]

def normal(t):
    return {
'T26':'''@Autowired private BookDao bookDao;\n@Autowired private BookDtoDoMapper mapper;\n\n@Override\npublic BookResponseDto createBook(BookCreateRequestDto request) {\n    BookDO book = mapper.toDO(request);\n    return mapper.toResponse(bookDao.save(book));\n}''',
'T27':'''@Autowired private BookDao bookDao;\n@Autowired private BookDtoDoMapper mapper;\n\n@Override\npublic BookResponseDto readBook(Long id) {\n    BookDO book = bookDao.findById(id).orElse(null);\n    return mapper.toResponse(book);\n}''',
'T28':'''@Autowired private BookDao bookDao;\n@Autowired private BookDtoDoMapper mapper;\n\n@Override\npublic BookResponseDto updateBook(Long id, BookUpdateRequestDto request) {\n    BookDO book = bookDao.findById(id).orElse(null);\n    mapper.applyUpdate(request, book);\n    return mapper.toResponse(bookDao.save(book));\n}''',
'T29':'''@Autowired private BookDao bookDao;\n@Autowired private BookDtoDoMapper mapper;\n\n@Override\npublic BookResponseDto deleteBook(Long id) {\n    BookDO book = bookDao.findById(id).orElse(null);\n    book.setActive(false);\n    return mapper.toResponse(bookDao.save(book));\n}''',
'T30':'''@Autowired private BookDao bookDao;\n@Autowired private BookDtoDoMapper mapper;\n\n@Override\npublic List<BookResponseDto> searchBook(String text) {\n    String normalized = text.trim();\n    return bookDao.search(normalized).stream().map(mapper::toResponse).toList();\n}'''}[t]

def corrected(t):
    return {
'T26':'''@Override\npublic BookResponseDto createBook(BookCreateRequestDto request) {\n    if (request == null) throw new IllegalArgumentException("Book request is required");\n    String isbn = request.getIsbn().trim();\n    bookDao.findByNormalizedBusinessKey(isbn).ifPresent(existing -> { throw new IllegalStateException("ISBN already exists"); });\n    request.setIsbn(isbn);\n    request.setTitle(request.getTitle().trim());\n    return mapper.toResponse(bookDao.save(mapper.toDO(request)));\n}''',
'T27':'''@Override\npublic BookResponseDto readBook(Long id) {\n    if (id == null) throw new IllegalArgumentException("Book ID is required");\n    BookDO book = bookDao.findCurrentById(id).orElseThrow(() -> new IllegalArgumentException("Book not found"));\n    return mapper.toResponse(book);\n}''',
'T28':'''@Override\npublic BookResponseDto updateBook(Long id, BookUpdateRequestDto request) {\n    if (id == null || request == null) throw new IllegalArgumentException("Book ID and request are required");\n    BookDO book = bookDao.findCurrentById(id).orElseThrow(() -> new IllegalArgumentException("Book not found"));\n    String isbn = request.getIsbn().trim();\n    bookDao.findByNormalizedBusinessKey(isbn).filter(other -> !other.getBookId().equals(id)).ifPresent(other -> { throw new IllegalStateException("ISBN already exists"); });\n    request.setIsbn(isbn); request.setTitle(request.getTitle().trim()); mapper.applyUpdate(request, book);\n    return mapper.toResponse(bookDao.save(book));\n}''',
'T29':'''@Override\npublic BookResponseDto deleteBook(Long id) {\n    if (id == null) throw new IllegalArgumentException("Book ID is required");\n    BookDO book = bookDao.findCurrentById(id).orElseThrow(() -> new IllegalArgumentException("Book not found"));\n    if (bookDao.countDependencies(id) > 0) throw new IllegalStateException("Book has active Book Copies");\n    book.setActive(false);\n    return mapper.toResponse(bookDao.save(book));\n}''',
'T30':'''@Override\npublic List<BookResponseDto> searchBook(String text) {\n    if (text == null) throw new IllegalArgumentException("Search text is required");\n    String normalized = text.trim();\n    if (normalized.isBlank()) throw new IllegalArgumentException("Search text is required");\n    return bookDao.searchIgnoreCase(normalized).stream().map(mapper::toResponse).toList();\n}'''}[t]

def observe(t):
    return {'T26':'service.createBook(null);','T27':'when(bookDao.findById(999L)).thenReturn(Optional.empty());\nservice.readBook(999L);','T28':'when(bookDao.findById(2L)).thenReturn(Optional.of(existing));\nservice.updateBook(2L, null);','T29':'when(bookDao.findById(999L)).thenReturn(Optional.empty());\nservice.deleteBook(999L);','T30':'service.searchBook(null);'}[t]

def itest(t):
    return {
'T26':'''@SpringBootTest\nclass CreateBookLocalIntegrationTest {\n @Autowired CreateBookService service;\n @Test void duplicateIsbnIsRejected() {\n   assertThrows(IllegalStateException.class, () -> service.createBook(request("9780132350884", "Another Clean Code", 1L,1L,1L)));\n }\n}''',
'T27':'''@SpringBootTest\nclass ReadBookLocalIntegrationTest {\n @Autowired ReadBookService service;\n @Test void activeBookIsRead() { assertEquals("9780134685991", service.readBook(2L).getIsbn()); }\n @Test void inactiveBookIsNotExposed() { assertThrows(IllegalArgumentException.class, () -> service.readBook(3L)); }\n}''',
'T28':'''@SpringBootTest\nclass UpdateBookLocalIntegrationTest {\n @Autowired UpdateBookService service;\n @Test void anotherBooksIsbnIsRejected() {\n   assertThrows(IllegalStateException.class, () -> service.updateBook(2L, request("9780132350884", "Effective Java", 2L,1L,2L)));\n }\n}''',
'T29':'''@SpringBootTest\nclass DeleteBookLocalIntegrationTest {\n @Autowired DeleteBookService service;\n @Test void bookWithOnlyWithdrawnCopyCanDeactivate() { assertFalse(service.deleteBook(2L).getActive()); }\n @Test void bookWithActiveCopiesCannotDeactivate() { assertThrows(IllegalStateException.class, () -> service.deleteBook(1L)); }\n}''',
'T30':'''@SpringBootTest\nclass SearchBookLocalIntegrationTest {\n @Autowired SearchBookService service;\n @Test void lowercaseTitleFindsStoredMixedCaseTitle() {\n   assertTrue(service.searchBook("clean code").stream().anyMatch(b -> "Clean Code".equals(b.getTitle())));\n }\n @Test void noMatchIsSuccessfulEmptyList() { assertTrue(service.searchBook("NO-SUCH-BOOK").isEmpty()); }\n}'''}[t]

def initial(t,d):
    doc=Document(); setup(doc); title(doc,f'{t} - {d["op"]} - Initial API Contract',f'{d["student"]} | {CONTRACT} | Book Module')
    doc.add_heading('1. Fixed Assignment',1); table(doc,[('Track',t),('Student',d['student']),('Operation',d['op']),('Service Code',d['code']),('Endpoint',d['endpoint']),('Service Interface',d['service']),('Method',sig(t)),('Success responseCode','00'),('Success Message',d['success'])])
    doc.add_heading('2. JSON Response Envelope',1); code(doc,f'{{\n  "serviceCode": "{d["code"]}",\n  "responseCode": "00",\n  "message": "{d["success"]}",\n  "data": {{ }}\n}}')
    doc.add_heading('3. Book Fields',1); table(doc,[('bookId','Database identifier; response only.'),('isbn','Book business key.'),('title','Book title.'),('authorId','Existing Author ID.'),('categoryId','Existing Category ID.'),('publisherId','Existing Publisher ID.'),('active','Logical active state; response only.')])
    doc.add_heading('4. Training Data',1); table(doc,[('Book 1','9780132350884 / Clean Code / ACTIVE'),('Book 2','9780134685991 / Effective Java / ACTIVE'),('Book 3','9780000000003 / Archived Programming Book / INACTIVE'),('Copy 1','ACC-0001 -> Book 1 / AVAILABLE'),('Copy 2','ACC-0002 -> Book 1 / ISSUED'),('Copy 3','ACC-0003 -> Book 2 / WITHDRAWN')])
    doc.add_heading('5. Ownership',1); table(doc,[('Presenter','REST Controller, DTO/DO, Mapper, base DAO, Flyway, JSON envelope and Thymeleaf framework.'),('Student',f'{d["impl"]}, Unit/Integration tests and {d["ui"]}.')]); note(doc,'The Initial Contract gives the successful starting behavior. Follow the Progressive Development Guide in order.')
    return doc

def progressive(t,d):
    doc=Document(); setup(doc); title(doc,f'{d["student"]} - {d["op"]} Progressive Testing Guide',f'{t} | Book Module | {CONTRACT}')
    doc.add_paragraph('Follow the sequence in order: prove invocation, implement the normal successful path, try the next realistic condition before changing the code, understand the result, correct the service, inspect coverage, test with real PostgreSQL, repeat with Testcontainers, and then build the assigned frontend.')
    doc.add_heading('1. Assignment and Flow',1); table(doc,[('Backend Assignment',d['op']),('Service',d['service']),('Implementation',d['impl']),('REST Endpoint',d['endpoint']),('Service Code',d['code']),('Frontend Assignment',d['ui']),('REST Controller','Presenter-owned')]); table(doc,[('Browser / REST client','Sends HTTP request.'),('Spring Web + Jackson','Converts JSON to supplied request DTO.'),('REST Controller','Calls frozen service interface.'),('Service','Student business logic.'),('Mapper','DTO <-> DO.'),('DAO','JPA persistence.'),('PostgreSQL','Real database.')]); note(doc,'The service receives Java DTO values, not raw JSON.')
    doc.add_heading('2. Create the Empty Service',1); doc.add_paragraph('The complete Spring chain has several components. The empty service removes that chain and proves one component can be invoked directly before database and HTTP concerns are involved. Mockito represents dependencies during Unit Testing.'); code(doc,f'@Service\npublic class {d["impl"]} implements {d["service"]} {{\n @Override public {sig(t)} {{ ' + ('return java.util.Collections.emptyList();' if t=='T30' else 'return null;') + ' }\n}')
    doc.add_heading('3. Invoke the Empty Service with JUnit',1); code(doc,f'@ExtendWith(MockitoExtension.class)\nclass {d["impl"]}Test {{\n @InjectMocks private {d["impl"]} service;\n @Test void shouldInvokeService() {{ /* invoke the assigned method */ }}\n}}'); bullets(doc,['Run only this test from Eclipse: Run As -> JUnit Test.','Add normal business logic only after this invocation test is green.'])
    doc.add_heading('4. Implement the Normal Successful Path',1); code(doc,normal(t)); doc.add_paragraph('Add one positive Mockito Unit Test for a valid request. Keep that test permanently as a regression test.')
    doc.add_heading('5. Try the Next Realistic Condition',1); doc.add_paragraph(d['unit']); doc.add_paragraph('Do not change the service before running this test. Observe the current behavior first.'); code(doc,observe(t)); doc.add_heading('5.1 Understand NullPointerException',2); doc.add_paragraph('A NullPointerException means a reference is null but code tries to call a method or access state through it. Read the first stack-trace line belonging to the service. Handle the application condition before dereferencing the null value instead of catching NullPointerException everywhere.')
    doc.add_heading('6. Improve the Service',1); code(doc,corrected(t)); doc.add_paragraph('Rerun the new test and the original positive test. The correction must preserve the earlier successful behavior.')
    doc.add_heading('7. JaCoCo Coverage',1); code(doc,'mvnw.cmd clean test'); table(doc,[('Green','Executed by tests.'),('Yellow','Partially covered branch.'),('Red','Not executed; there is still room for uncertainty.')]); doc.add_paragraph('Open target/site/jacoco/index.html and inspect the Book service implementation.')
    doc.add_heading('8. Why Integration Testing is Required',1); doc.add_paragraph('Mockito does not prove real ISBN duplication, active/inactive rows, Book Copy status, Flyway data, JPA mappings or PostgreSQL comparison behavior. Integration Testing joins the real components.'); table(doc,[('Unit Test','Service + Mockito mocks.'),('Local Integration','Spring + Mapper + DAO + local PostgreSQL.'),('Testcontainers','Spring + Mapper + DAO + temporary PostgreSQL.')])
    doc.add_heading('9. PostgreSQL JDBC and Commented MySQL Reference',1); code(doc,'spring.datasource.url=jdbc:postgresql://localhost:5432/library_full_stack\nspring.datasource.username=postgres\nspring.datasource.password=postgres\nspring.datasource.driver-class-name=org.postgresql.Driver'); code(doc,'<dependency>\n <groupId>org.postgresql</groupId>\n <artifactId>postgresql</artifactId>\n <scope>runtime</scope>\n</dependency>\n\n<!-- MySQL learning reference only\n<dependency>\n <groupId>com.mysql</groupId>\n <artifactId>mysql-connector-j</artifactId>\n <scope>runtime</scope>\n</dependency>\n-->'); code(doc,'# MySQL reference only - keep commented\n# spring.datasource.url=jdbc:mysql://localhost:3306/library_full_stack\n# spring.datasource.driver-class-name=com.mysql.cj.jdbc.Driver')
    doc.add_heading('10. Inspect the Training Data with DBeaver',1); code(doc,'SELECT book_id, isbn, title, is_active FROM tbl_book ORDER BY book_id;\n\nSELECT book_copy_id, accession_number, fk_book, status FROM tbl_book_copy ORDER BY book_copy_id;'); doc.add_paragraph(d['integration'])
    doc.add_heading('11. Complete Local Integration Test',1); code(doc,itest(t))
    if t=='T30': doc.add_heading('11.1 Correct the PostgreSQL Search',2); code(doc,'@Transactional(readOnly = true)\npublic List<BookDO> searchIgnoreCase(String text) {\n String pattern = "%" + text.toLowerCase() + "%";\n return entityManager.createQuery("select b from BookDO b where b.active = true and (lower(b.isbn) like :pattern or lower(b.title) like :pattern) order by b.bookId", BookDO.class).setParameter("pattern", pattern).getResultList();\n}')
    doc.add_heading('12. PostgreSQL Testcontainers',1); code(doc,'@Container\n@ServiceConnection\nstatic PostgreSQLContainer postgres =\n new PostgreSQLContainer(DockerImageName.parse("postgres:18"));'); doc.add_paragraph('Add @Testcontainers to the Integration Test class and repeat the same database cases. Flyway recreates deterministic training data in the temporary container.')
    doc.add_heading('13. Final Maven Regression',1); code(doc,'mvnw.cmd clean test'); bullets(doc,['Focused Unit Test green.','Original successful Unit Test still green.','Local PostgreSQL Integration Test green.','PostgreSQL Testcontainers test green.','JaCoCo reviewed.'])
    doc.add_heading('14. Student-owned Files',1); table(doc,[('Service',d['impl']+'.java'),('Tests','Track-specific Unit and Integration Test classes.'),('Frontend',d['ui']),('Presenter-owned','REST Controller, DTO/DO, Mapper, base DAO, Flyway and Thymeleaf framework.')]);
    if t=='T30': note(doc,'T30 owns only the case-independent Book search DAO method named in this guide.')
    doc.add_heading('15. Frontend Assignment',1); doc.add_paragraph(f'{d["student"]} implements the {d["ui"]}. First prove main.js and the button handler, then build the page in one JavaScript file. Test against Presenter STUB/support before partner integration.'); code(doc,'<script type="module" src="/src/main.js"></script>\n\nconsole.log("main.js loaded");\ndocument.querySelector("#actionButton").addEventListener("click", () => console.log("button reached JavaScript"));')
    doc.add_heading('16. HTTP Transport Handling',1); code(doc,'if (response.status === 404) { showMessage("NOT Found"); return; }\nif (!response.ok) { showMessage("Error Encountered, Please contact Administrator"); return; }\nconst body = await response.json();\nshowMessage(body.message);')
    doc.add_heading('17. Refactor After the Page Works',1); table(doc,[('src/main.js','Starts page and connects modules.'),('src/api/bookApi.js','Contains REST calls.'),('src/forms/bookForm.js','Reads form values.'),('src/views/bookView.js','Renders messages and records.')]); doc.add_paragraph('Keep the same /rest URL when switching from Presenter STUB/support to the real partner implementation.')
    doc.add_heading('18. Completion Checklist',1); bullets(doc,['Empty service invoked.','Normal path Unit Test retained.','Realistic next condition observed before correction.','Service corrected.','JaCoCo checked.','Local PostgreSQL test completed.','Testcontainers completed.','Single-file frontend proved.','Frontend refactored only afterward.','Partner integration uses unchanged URL.'])
    return doc

def updated(t,d):
    doc=Document(); setup(doc); title(doc,f'{t} - {d["op"]} - Updated API Contract',f'{CONTRACT} | Book Module'); doc.add_heading('1. Fixed Endpoint',1); table(doc,[('Track',t),('Service Code',d['code']),('Endpoint',d['endpoint']),('Controlled application result','HTTP 200')]); doc.add_heading('2. Response Envelope',1); code(doc,'{\n  "serviceCode":"XX",\n  "responseCode":"00",\n  "message":"...",\n  "data":{}\n}')
    rows=[('00',d['success'])]
    if t=='T26': rows += [('01','Invalid / missing Book request'),('03','ISBN already exists')]
    elif t=='T27': rows += [('01','Book ID is required'),('02','Book not found / inactive Book not exposed')]
    elif t=='T28': rows += [('01','Invalid Book ID/request'),('02','Book not found'),('03','ISBN belongs to another Book')]
    elif t=='T29': rows += [('01','Book ID is required'),('02','Book not found'),('05','Book has non-withdrawn Book Copy dependencies')]
    else: rows += [('01','Search text is required or blank')]
    doc.add_heading('3. Response Codes',1); table(doc,rows,('responseCode','Meaning')); doc.add_heading('4. Frontend Transport Errors',1); table(doc,[('HTTP 404','NOT Found'),('Other non-200','Error Encountered, Please contact Administrator')]); note(doc,'A valid Search with zero matches returns responseCode 00 and an empty list.' if t=='T30' else 'Controlled application results use the frozen JSON response envelope and HTTP 200.'); return doc

for t,d in TRACKS.items():
    folder=OUT/f'{t}_{d["op"].replace(" ","_")}'
    folder.mkdir(parents=True,exist_ok=True)
    initial(t,d).save(folder/f'{t}_01_Initial_API_Contract.docx')
    progressive(t,d).save(folder/f'{t}_02_Progressive_Development_Guide.docx')
    updated(t,d).save(folder/f'{t}_03_Updated_API_Contract.docx')

print('Generated Book guides under', OUT)
