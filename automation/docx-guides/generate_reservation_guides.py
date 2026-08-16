from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import textwrap

CONTRACT = "LIBRARY-FULLSTACK-V1"
ROOT = Path("Documents/Student-Guides/Reservation")
ROOT.mkdir(parents=True, exist_ok=True)

TRACKS = {
    "T51": {
        "student": "Student 51", "op": "Create Reservation", "code": "51",
        "service": "CreateReservationService", "impl": "CreateReservationServiceImpl",
        "endpoint": "POST /rest/reservations", "success": "Reservation Created Successfully",
        "frontend": "Update Reservation UI", "partner": "Student 53 - Update Reservation",
        "unit": "After the valid Create test is green, call createReservation(null). Run the test before adding validation.",
        "integration": "The PostgreSQL training data already contains RES-0001. Try to create another Reservation with the same reservationNumber.",
    },
    "T52": {
        "student": "Student 52", "op": "Read Reservation", "code": "52",
        "service": "ReadReservationService", "impl": "ReadReservationServiceImpl",
        "endpoint": "GET /rest/reservations/{id}", "success": "Reservation Read Successfully",
        "frontend": "Cancel Reservation UI", "partner": "Student 54 - Cancel Reservation",
        "unit": "After the valid Read test is green, request Reservation ID 999. Configure ReservationDao to return Optional.empty() and observe the current behavior.",
        "integration": "Reservation ID 2 physically exists as RES-0002 but its status is CANCELLED. Normal Read must not expose this non-current Reservation.",
    },
    "T53": {
        "student": "Student 53", "op": "Update Reservation", "code": "53",
        "service": "UpdateReservationService", "impl": "UpdateReservationServiceImpl",
        "endpoint": "PUT /rest/reservations/{id}", "success": "Reservation Updated Successfully",
        "frontend": "Search Reservation UI", "partner": "Student 55 - Search Reservation",
        "unit": "After the valid Update test is green, call updateReservation(1L, null). Run the test before adding validation.",
        "integration": "Reservation ID 1 owns RES-0001 and Reservation ID 2 owns RES-0002. Try to update ID 2 to RES-0001. Exclude the current row from duplicate detection.",
    },
    "T54": {
        "student": "Student 54", "op": "Cancel Reservation", "code": "54",
        "service": "DeleteReservationService", "impl": "DeleteReservationServiceImpl",
        "endpoint": "DELETE /rest/reservations/{id}", "success": "Reservation Cancelled Successfully",
        "frontend": "Create Reservation UI", "partner": "Student 51 - Create Reservation",
        "unit": "After the normal Cancel path is green, request Reservation ID 999. Run the test before adding not-found handling.",
        "integration": "RES-0004 is ACTIVE but a non-CANCELLED Book Issue is already linked to it. The persisted relationship must prevent cancellation.",
    },
    "T55": {
        "student": "Student 55", "op": "Search Reservation", "code": "55",
        "service": "SearchReservationService", "impl": "SearchReservationServiceImpl",
        "endpoint": "GET /rest/reservations/search?text={text}", "success": "Reservation Search Completed Successfully",
        "frontend": "Reservation List / Search Results UI", "partner": "Presenter Reservation List Support",
        "unit": "After normal Search is green, call searchReservation(null). The normal first implementation calls trim(), so run the test before adding validation.",
        "integration": "Search for res-0001 while PostgreSQL stores RES-0001. The completed search must be case-independent and ACTIVE-only.",
    },
}


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), fill)
    pr.append(sh)


def setup(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.62)
        section.bottom_margin = Inches(0.62)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(10.2)
    for name, size in (("Title", 21), ("Heading 1", 15), ("Heading 2", 12.7), ("Heading 3", 11.1)):
        doc.styles[name].font.name = "Aptos"
        doc.styles[name].font.size = Pt(size)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.bold = True


def add_table(doc, rows, headers=("Item", "Meaning / Value")):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text, t.cell(0, 1).text = headers
    shade(t.cell(0, 0), "D9EAF7")
    shade(t.cell(0, 1), "D9EAF7")
    for left, right in rows:
        cells = t.add_row().cells
        cells[0].text = str(left)
        cells[1].text = str(right)


def add_code(doc, source):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(textwrap.dedent(source).strip())
    r.font.name = "Consolas"
    r.font.size = Pt(8.05)
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), "F3F4F6")
    p._p.get_or_add_pPr().append(sh)


def add_note(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    shade(t.cell(0, 0), "FFF7D6")
    p = t.cell(0, 0).paragraphs[0]
    r = p.add_run("Important: ")
    r.bold = True
    p.add_run(text)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def signature(track):
    return {
        "T51": "ReservationResponseDto createReservation(ReservationCreateRequestDto request)",
        "T52": "ReservationResponseDto readReservation(Long id)",
        "T53": "ReservationResponseDto updateReservation(Long id, ReservationUpdateRequestDto request)",
        "T54": "ReservationResponseDto deleteReservation(Long id)",
        "T55": "List<ReservationResponseDto> searchReservation(String text)",
    }[track]


def normal_code(track):
    return {
        "T51": """@Autowired private ReservationDao reservationDao;\n@Autowired private ReservationDtoDoMapper mapper;\n\n@Override\npublic ReservationResponseDto createReservation(ReservationCreateRequestDto request) {\n    ReservationDO dataObject = mapper.toDO(request);\n    return mapper.toResponse(reservationDao.save(dataObject));\n}""",
        "T52": """@Autowired private ReservationDao reservationDao;\n@Autowired private ReservationDtoDoMapper mapper;\n\n@Override\npublic ReservationResponseDto readReservation(Long id) {\n    ReservationDO reservation = reservationDao.findById(id).orElse(null);\n    return mapper.toResponse(reservation);\n}""",
        "T53": """@Autowired private ReservationDao reservationDao;\n@Autowired private ReservationDtoDoMapper mapper;\n\n@Override\npublic ReservationResponseDto updateReservation(Long id, ReservationUpdateRequestDto request) {\n    ReservationDO reservation = reservationDao.findById(id).orElse(null);\n    mapper.applyUpdate(request, reservation);\n    return mapper.toResponse(reservationDao.save(reservation));\n}""",
        "T54": """@Autowired private ReservationDao reservationDao;\n@Autowired private ReservationDtoDoMapper mapper;\n\n@Override\npublic ReservationResponseDto deleteReservation(Long id) {\n    ReservationDO reservation = reservationDao.findById(id).orElse(null);\n    reservation.setStatus(\"CANCELLED\");\n    return mapper.toResponse(reservationDao.save(reservation));\n}""",
        "T55": """@Autowired private ReservationDao reservationDao;\n@Autowired private ReservationDtoDoMapper mapper;\n\n@Override\npublic List<ReservationResponseDto> searchReservation(String text) {\n    String normalized = text.trim();\n    return reservationDao.search(normalized).stream().map(mapper::toResponse).toList();\n}""",
    }[track]


def observe_code(track):
    return {
        "T51": "service.createReservation(null);",
        "T52": "when(reservationDao.findById(999L)).thenReturn(Optional.empty());\nservice.readReservation(999L);",
        "T53": "when(reservationDao.findById(1L)).thenReturn(Optional.of(existing));\nservice.updateReservation(1L, null);",
        "T54": "when(reservationDao.findById(999L)).thenReturn(Optional.empty());\nservice.deleteReservation(999L);",
        "T55": "service.searchReservation(null);",
    }[track]


def corrected_code(track):
    return {
        "T51": """@Override\npublic ReservationResponseDto createReservation(ReservationCreateRequestDto request) {\n    if (request == null) throw new IllegalArgumentException(\"Reservation request is required\");\n    String number = request.getReservationNumber().trim();\n    reservationDao.findByNormalizedBusinessKey(number).ifPresent(existing -> {\n        throw new IllegalStateException(\"Reservation number already exists\");\n    });\n    request.setReservationNumber(number);\n    request.setStatus(request.getStatus().trim());\n    return mapper.toResponse(reservationDao.save(mapper.toDO(request)));\n}""",
        "T52": """@Override\npublic ReservationResponseDto readReservation(Long id) {\n    if (id == null) throw new IllegalArgumentException(\"Reservation ID is required\");\n    ReservationDO reservation = reservationDao.findCurrentById(id)\n        .orElseThrow(() -> new IllegalArgumentException(\"Reservation not found\"));\n    return mapper.toResponse(reservation);\n}""",
        "T53": """@Override\npublic ReservationResponseDto updateReservation(Long id, ReservationUpdateRequestDto request) {\n    if (id == null || request == null) throw new IllegalArgumentException(\"Reservation ID and request are required\");\n    ReservationDO reservation = reservationDao.findById(id)\n        .orElseThrow(() -> new IllegalArgumentException(\"Reservation not found\"));\n    String number = request.getReservationNumber().trim();\n    reservationDao.findByNormalizedBusinessKey(number)\n        .filter(other -> !other.getReservationId().equals(id))\n        .ifPresent(other -> { throw new IllegalStateException(\"Reservation number already exists\"); });\n    request.setReservationNumber(number);\n    request.setStatus(request.getStatus().trim());\n    mapper.applyUpdate(request, reservation);\n    return mapper.toResponse(reservationDao.save(reservation));\n}""",
        "T54": """@Override\npublic ReservationResponseDto deleteReservation(Long id) {\n    if (id == null) throw new IllegalArgumentException(\"Reservation ID is required\");\n    ReservationDO reservation = reservationDao.findCurrentById(id)\n        .orElseThrow(() -> new IllegalArgumentException(\"Active Reservation not found\"));\n    if (reservationDao.countDependencies(id) > 0) {\n        throw new IllegalStateException(\"Reservation already has a Book Issue\");\n    }\n    reservation.setStatus(\"CANCELLED\");\n    return mapper.toResponse(reservationDao.save(reservation));\n}""",
        "T55": """@Override\npublic List<ReservationResponseDto> searchReservation(String text) {\n    if (text == null) throw new IllegalArgumentException(\"Search text is required\");\n    String normalized = text.trim();\n    if (normalized.isBlank()) throw new IllegalArgumentException(\"Search text is required\");\n    return reservationDao.searchIgnoreCaseCurrent(normalized).stream().map(mapper::toResponse).toList();\n}""",
    }[track]


def integration_code(track):
    return {
        "T51": """@SpringBootTest\nclass CreateReservationLocalIntegrationTest {\n    @Autowired private CreateReservationService service;\n\n    @Test\n    void shouldCreateNewReservationNumber() {\n        ReservationResponseDto result = service.createReservation(\n            request(\"RES-NEW-001\", 4L, 4L, LocalDate.of(2026, 8, 16), \"ACTIVE\"));\n        assertEquals(\"RES-NEW-001\", result.getReservationNumber());\n    }\n\n    @Test\n    void shouldRejectPersistedDuplicateReservationNumber() {\n        assertThrows(IllegalStateException.class, () ->\n            service.createReservation(\n                request(\"RES-0001\", 4L, 4L, LocalDate.of(2026, 8, 16), \"ACTIVE\")));\n    }\n}""",
        "T52": """@SpringBootTest\nclass ReadReservationLocalIntegrationTest {\n    @Autowired private ReadReservationService service;\n\n    @Test void shouldReadActiveReservation() {\n        ReservationResponseDto result = service.readReservation(1L);\n        assertEquals(\"RES-0001\", result.getReservationNumber());\n        assertEquals(\"ACTIVE\", result.getStatus());\n    }\n\n    @Test void normalReadShouldNotExposeCancelledReservation() {\n        assertThrows(IllegalArgumentException.class, () -> service.readReservation(2L));\n    }\n}""",
        "T53": """@SpringBootTest\nclass UpdateReservationLocalIntegrationTest {\n    @Autowired private UpdateReservationService service;\n\n    @Test void shouldAllowCurrentReservationToKeepOwnNumber() {\n        ReservationResponseDto result = service.updateReservation(1L,\n            request(\"RES-0001\", 2L, 1L, LocalDate.of(2026, 8, 1), \"ACTIVE\"));\n        assertEquals(\"RES-0001\", result.getReservationNumber());\n    }\n\n    @Test void shouldRejectNumberOwnedByAnotherReservation() {\n        assertThrows(IllegalStateException.class, () -> service.updateReservation(2L,\n            request(\"RES-0001\", 4L, 2L, LocalDate.of(2026, 8, 2), \"CANCELLED\")));\n    }\n}""",
        "T54": """@SpringBootTest\nclass CancelReservationLocalIntegrationTest {\n    @Autowired private DeleteReservationService service;\n\n    @Test void shouldCancelActiveReservationWithoutIssue() {\n        ReservationResponseDto result = service.deleteReservation(1L);\n        assertEquals(\"CANCELLED\", result.getStatus());\n    }\n\n    @Test void shouldNotCancelReservationAlreadyUsedByBookIssue() {\n        assertThrows(IllegalStateException.class, () -> service.deleteReservation(4L));\n    }\n}""",
        "T55": """@SpringBootTest\nclass SearchReservationLocalIntegrationTest {\n    @Autowired private SearchReservationService service;\n\n    @Test void lowercaseSearchFindsStoredUppercaseNumber() {\n        assertTrue(service.searchReservation(\"res-0001\").stream()\n            .anyMatch(r -> \"RES-0001\".equals(r.getReservationNumber())));\n    }\n\n    @Test void searchDoesNotReturnCancelledReservation() {\n        assertTrue(service.searchReservation(\"res-0002\").isEmpty());\n    }\n\n    @Test void validNoMatchReturnsEmptyList() {\n        assertTrue(service.searchReservation(\"NO-SUCH-RESERVATION\").isEmpty());\n    }\n}""",
    }[track]


def initial_doc(track, data):
    doc = Document(); setup(doc)
    add_title(doc, f"{track} - {data['op']} - Initial API Contract", f"{data['student']} | {CONTRACT} | Reservation Module")
    doc.add_heading("1. Fixed Assignment", level=1)
    add_table(doc, [("Track", track), ("Student", data["student"]), ("Operation", data["op"]),
                    ("Service Code", data["code"]), ("Endpoint", data["endpoint"]),
                    ("Service Interface", data["service"]), ("Method", signature(track)),
                    ("Success responseCode", "00"), ("Success Message", data["success"])])
    doc.add_heading("2. Frozen JSON Response Envelope", level=1)
    add_code(doc, f'''{{\n  "serviceCode": "{data["code"]}",\n  "responseCode": "00",\n  "message": "{data["success"]}",\n  "data": {{ }}\n}}''')
    doc.add_heading("3. Reservation Fields", level=1)
    add_table(doc, [("reservationId", "Database identifier; response only."),
                    ("reservationNumber", "Business key such as RES-0001."),
                    ("membershipId", "Membership making the Reservation."),
                    ("bookId", "Book being reserved."),
                    ("reservedDate", "Date on which the Reservation was recorded."),
                    ("status", "Lifecycle value: ACTIVE, CANCELLED or FULFILLED.")])
    doc.add_heading("4. Training Data", level=1)
    add_table(doc, [("Reservation 1", "RES-0001 / ACTIVE"),
                    ("Reservation 2", "RES-0002 / CANCELLED"),
                    ("Reservation 3", "RES-0003 / FULFILLED / used by ISS-0003"),
                    ("Reservation 4", "RES-0004 / ACTIVE / linked to a non-CANCELLED Book Issue")])
    doc.add_heading("5. Ownership", level=1)
    add_table(doc, [("Presenter", "REST Controller, DTO/DO, Mapper, base DAO, Flyway, JSON response and Thymeleaf framework."),
                    ("Student", f"{data['impl']}, track tests and {data['frontend']}.")])
    add_note(doc, "Delete for Reservation means Cancel. The row is retained and status becomes CANCELLED.")
    return doc


def progressive_doc(track, data):
    doc = Document(); setup(doc)
    add_title(doc, f"{data['student']} - {data['op']} Progressive Testing Guide", f"{track} | Reservation Module | {CONTRACT}")
    doc.add_paragraph("Follow the approved progression: empty service, focused invocation, natural successful implementation, positive Unit Test, next realistic condition, observation, correction, regression, JaCoCo, local PostgreSQL, complete Integration Test, PostgreSQL Testcontainers and frontend integration.")
    doc.add_heading("1. Assignment and Flow", level=1)
    add_table(doc, [("Backend Assignment", data["op"]), ("Service", data["service"]),
                    ("Implementation", data["impl"]), ("REST Endpoint", data["endpoint"]),
                    ("Service Code", data["code"]), ("Frontend Assignment", data["frontend"]),
                    ("REST Controller", "Presenter-owned")])
    add_table(doc, [("Browser / REST client", "Sends HTTP request."),
                    ("Spring Web + Jackson", "Creates the request DTO for POST/PUT."),
                    ("REST Controller", "Calls the frozen service interface."),
                    ("Student Service", "Implements the business behavior."),
                    ("Mapper", "Converts DTO <-> ReservationDO."),
                    ("DAO", "Runs JPA persistence."),
                    ("PostgreSQL", "Stores Reservation and Book Issue relationships.")])
    add_note(doc, "The service receives Java DTO values, not raw JSON.")

    doc.add_heading("2. Create the Empty Service", level=1)
    doc.add_paragraph("The empty service proves the assigned component can be reached without the entire Spring request chain. Mockito is later used to provide controlled DAO and Mapper dependencies during Unit Testing.")
    ret = "return java.util.Collections.emptyList();" if track == "T55" else "return null;"
    add_code(doc, f'''@Service\npublic class {data["impl"]} implements {data["service"]} {{\n    @Override\n    public {signature(track)} {{\n        {ret}\n    }}\n}}''')
    doc.add_heading("3. Invoke the Empty Service", level=1)
    add_code(doc, f'''@ExtendWith(MockitoExtension.class)\nclass {data["impl"]}Test {{\n    @InjectMocks private {data["impl"]} service;\n    @Test void shouldInvokeService() {{\n        // Invoke the assigned method.\n    }}\n}}''')
    bullets(doc, ["Run the focused test from Eclipse using Run As -> JUnit Test.", "Add normal business logic only after the invocation test is green."])

    doc.add_heading("4. Implement the Natural Successful Path", level=1)
    add_code(doc, normal_code(track))
    doc.add_paragraph("Add one positive Mockito Unit Test for a valid Reservation and keep it as a regression test.")

    doc.add_heading("5. Try the Next Realistic Condition", level=1)
    doc.add_paragraph(data["unit"])
    doc.add_paragraph("Do not change the service before running the test. Observe what the current implementation does.")
    add_code(doc, observe_code(track))
    doc.add_heading("5.1 Understand the failure", level=2)
    doc.add_paragraph("Read the exception type and the first stack-trace line from your service. A NullPointerException means a null reference was used as though an object existed. Model the application condition before the value is dereferenced.")

    doc.add_heading("6. Improve the Service", level=1)
    add_code(doc, corrected_code(track))
    doc.add_paragraph("Rerun the new test and the original successful test. Both must pass.")

    doc.add_heading("7. JaCoCo Coverage", level=1)
    add_code(doc, "mvnw.cmd clean test")
    add_table(doc, [("Green", "Executed by tests."), ("Yellow", "Partially covered branch."),
                    ("Red", "Not executed; there is still room for uncertainty.")])
    doc.add_paragraph("Open target/site/jacoco/index.html and inspect the assigned Reservation service.")

    doc.add_heading("8. Integration Testing", level=1)
    doc.add_paragraph("Mockito cannot prove real reservation numbers, ACTIVE/CANCELLED lifecycle state, JPA mappings or the persisted Reservation-to-Book-Issue relationship. Integration Testing joins the real components.")
    add_table(doc, [("Unit Test", "Service + Mockito mocks."),
                    ("Local Integration Test", "Spring + Mapper + DAO + local PostgreSQL."),
                    ("Testcontainers", "Same path with temporary PostgreSQL 18.")])

    doc.add_heading("9. PostgreSQL JDBC and MySQL Reference", level=1)
    add_code(doc, '''spring.datasource.url=jdbc:postgresql://localhost:5432/library_full_stack\nspring.datasource.username=postgres\nspring.datasource.password=postgres\nspring.datasource.driver-class-name=org.postgresql.Driver''')
    add_code(doc, '''<dependency>\n    <groupId>org.postgresql</groupId>\n    <artifactId>postgresql</artifactId>\n    <scope>runtime</scope>\n</dependency>\n\n<!-- MySQL learning reference only - keep commented\n<dependency>\n    <groupId>com.mysql</groupId>\n    <artifactId>mysql-connector-j</artifactId>\n    <scope>runtime</scope>\n</dependency>\n-->''')
    add_code(doc, '''# MySQL reference only - DO NOT activate\n# spring.datasource.url=jdbc:mysql://localhost:3306/library_full_stack\n# spring.datasource.driver-class-name=com.mysql.cj.jdbc.Driver''')

    doc.add_heading("10. Inspect the Real Database with DBeaver", level=1)
    add_code(doc, '''SELECT reservation_id, reservation_number, fk_membership, fk_book, reserved_date, status\nFROM tbl_reservation\nORDER BY reservation_id;\n\nSELECT book_issue_id, issue_number, fk_reservation, status\nFROM tbl_book_issue\nWHERE fk_reservation IS NOT NULL\nORDER BY book_issue_id;''')
    doc.add_paragraph(data["integration"])
    if track == "T54":
        add_note(doc, "RES-0004 is ACTIVE while a non-CANCELLED Book Issue is already linked to it. A multi-step workflow can temporarily contain this state; the persisted relationship must still prevent cancellation.")

    doc.add_heading("11. Complete Local PostgreSQL Integration Test", level=1)
    add_code(doc, integration_code(track))
    if track == "T55":
        doc.add_heading("11.1 Correct the PostgreSQL Search", level=2)
        add_code(doc, '''@Transactional(readOnly = true)\npublic List<ReservationDO> searchIgnoreCaseCurrent(String text) {\n    String pattern = "%" + text.toLowerCase() + "%";\n    return entityManager.createQuery(\n        "select r from ReservationDO r where r.status = 'ACTIVE' " +\n        "and lower(r.reservationNumber) like :pattern order by r.reservationId",\n        ReservationDO.class)\n        .setParameter("pattern", pattern)\n        .getResultList();\n}''')

    doc.add_heading("12. PostgreSQL Testcontainers", level=1)
    add_code(doc, '''@Container\n@ServiceConnection\nstatic PostgreSQLContainer postgres =\n    new PostgreSQLContainer(\n        DockerImageName.parse("postgres:18"));''')
    doc.add_paragraph("Add @Testcontainers to the Integration Test class and repeat the same database cases. Flyway recreates the deterministic Reservation data in the temporary container.")

    doc.add_heading("13. Final Regression", level=1)
    add_code(doc, "mvnw.cmd clean test")
    bullets(doc, ["Focused Unit Test green.", "Original positive Unit Test still green.",
                  "Local PostgreSQL Integration Test green.", "PostgreSQL Testcontainers test green.",
                  "JaCoCo reviewed."])

    doc.add_heading("14. Student-owned Files", level=1)
    add_table(doc, [("Service", data["impl"] + ".java"),
                    ("Tests", "Track-specific Unit and Integration Test classes."),
                    ("Frontend", data["frontend"]),
                    ("Presenter-owned", "REST Controller, DTO/DO, Mapper, base DAO, Flyway and Thymeleaf framework.")])
    if track == "T55":
        add_note(doc, "T55 owns only the case-independent current Reservation search DAO method named by this guide.")

    doc.add_heading("15. Frontend Assignment", level=1)
    doc.add_paragraph(f"{data['student']} implements the {data['frontend']}. First prove main.js and the button handler, then build a working page in one JavaScript file. Test against Presenter STUB/support before integrating {data['partner']}. Keep the same /rest URL.")
    add_code(doc, '''<script type="module" src="/src/main.js"></script>\n\nconsole.log("main.js loaded");\ndocument.querySelector("#actionButton")\n  .addEventListener("click", () => {\n    console.log("action button reached JavaScript");\n  });''')
    doc.add_heading("16. HTTP Transport Handling", level=1)
    add_code(doc, '''if (response.status === 404) {\n  showMessage("NOT Found");\n  return;\n}\nif (!response.ok) {\n  showMessage("Error Encountered, Please contact Administrator");\n  return;\n}\nconst body = await response.json();\nshowMessage(body.message);''')
    doc.add_heading("17. Refactor After the Page Works", level=1)
    add_table(doc, [("src/main.js", "Starts the page."),
                    ("src/api/reservationApi.js", "Contains Reservation REST calls."),
                    ("src/forms/reservationForm.js", "Reads page values."),
                    ("src/views/reservationView.js", "Renders records and messages.")])
    bullets(doc, ["Presenter STUB/support works first.", "Partner integration uses the unchanged URL.",
                  "Assigned files are committed and pushed."])
    return doc


def updated_doc(track, data):
    doc = Document(); setup(doc)
    add_title(doc, f"{track} - {data['op']} - Updated API Contract", f"{CONTRACT} | Reservation Module")
    doc.add_heading("1. Fixed Endpoint", level=1)
    add_table(doc, [("Track", track), ("Service Code", data["code"]), ("Endpoint", data["endpoint"]),
                    ("Controlled application result", "HTTP 200")])
    doc.add_heading("2. Frozen JSON Envelope", level=1)
    add_code(doc, '''{\n  "serviceCode": "XX",\n  "responseCode": "00",\n  "message": "...",\n  "data": {}\n}''')
    rows = [("00", data["success"])]
    if track == "T51": rows += [("01", "Invalid / missing Reservation request"), ("03", "Reservation number already exists")]
    elif track == "T52": rows += [("01", "Reservation ID is required"), ("02", "Reservation not found / non-ACTIVE row not exposed")]
    elif track == "T53": rows += [("01", "Invalid Reservation ID/request"), ("02", "Reservation not found"), ("03", "Reservation number belongs to another Reservation")]
    elif track == "T54": rows += [("01", "Reservation ID is required"), ("02", "Active Reservation not found"), ("05", "Reservation already has a non-CANCELLED Book Issue")]
    else: rows += [("01", "Search text is required or blank")]
    doc.add_heading("3. Response Codes", level=1)
    add_table(doc, rows, ("responseCode", "Meaning"))
    doc.add_heading("4. Frontend Transport Errors", level=1)
    add_table(doc, [("HTTP 404", "NOT Found"),
                    ("Other non-200", "Error Encountered, Please contact Administrator")])
    add_note(doc, "A valid Search with zero matches returns responseCode 00 and an empty list." if track == "T55" else "Controlled application results use the frozen JSON response envelope and HTTP 200.")
    return doc


for track, data in TRACKS.items():
    folder = ROOT / f"{track}_{data['op'].replace(' ', '_')}"
    folder.mkdir(parents=True, exist_ok=True)
    initial_doc(track, data).save(folder / f"{track}_01_Initial_API_Contract.docx")
    progressive_doc(track, data).save(folder / f"{track}_02_Progressive_Development_Guide.docx")
    updated_doc(track, data).save(folder / f"{track}_03_Updated_API_Contract.docx")

print("Generated Reservation T51-T55 guides in", ROOT)
