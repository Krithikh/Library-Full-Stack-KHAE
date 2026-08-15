from pathlib import Path
import textwrap
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('Documents/Student-Guides/Category')
ROOT.mkdir(parents=True, exist_ok=True)
CONTRACT = 'LIBRARY-FULLSTACK-V1'

TRACKS = {
'T16': dict(student='Student 16', op='Create Category', service='CreateCategoryService', impl='CreateCategoryServiceImpl', code='16', endpoint='POST /rest/categories', success='Category Created Successfully', ui='Update Category UI', discovery='After the successful Create test works, call createCategory(null). Run the test before adding validation.', integration='The training database already contains PROGRAMMING. Try to create another Category with categoryCode PROGRAMMING.'),
'T17': dict(student='Student 17', op='Read Category', service='ReadCategoryService', impl='ReadCategoryServiceImpl', code='17', endpoint='GET /rest/categories/{id}', success='Category Read Successfully', ui='Deactivate Category UI', discovery='After a successful Read works, request Category ID 999. Configure the mock DAO to return Optional.empty() and observe the current behavior.', integration='Category ID 3 physically exists as LEGACY but is inactive. Normal Read must not expose this inactive row.'),
'T18': dict(student='Student 18', op='Update Category', service='UpdateCategoryService', impl='UpdateCategoryServiceImpl', code='18', endpoint='PUT /rest/categories/{id}', success='Category Updated Successfully', ui='Search Category UI', discovery='After the successful Update test works, call updateCategory(2L, null). Run the test before changing the service.', integration='Category ID 1 owns PROGRAMMING and Category ID 2 owns DATABASE. Try to update ID 2 to PROGRAMMING. The current row must be excluded from duplicate detection.'),
'T19': dict(student='Student 19', op='Deactivate Category', service='DeleteCategoryService', impl='DeleteCategoryServiceImpl', code='19', endpoint='DELETE /rest/categories/{id}', success='Category Deactivated Successfully', ui='Create Category UI', discovery='After normal logical deactivation works, request Category ID 999 and observe the current behavior.', integration='Category PROGRAMMING is used by active Books. It must not be deactivated while those persisted dependencies exist.'),
'T20': dict(student='Student 20', op='Search Category', service='SearchCategoryService', impl='SearchCategoryServiceImpl', code='20', endpoint='GET /rest/categories/search?text={text}', success='Category Search Completed Successfully', ui='Category List / Search Results UI', discovery='After normal Search works, call searchCategory(null). The normal first implementation trims the text, so run the test before changing it.', integration='Search for programming when PostgreSQL stores PROGRAMMING and Programming. The completed search must be case-independent.'),
}

def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), fill); pr.append(sh)

def setup(doc):
    for s in doc.sections:
        s.top_margin=Inches(.62); s.bottom_margin=Inches(.62); s.left_margin=Inches(.72); s.right_margin=Inches(.72)
    doc.styles['Normal'].font.name='Aptos'; doc.styles['Normal'].font.size=Pt(10.2)
    for n,sz in [('Title',21),('Heading 1',15),('Heading 2',12.6),('Heading 3',11.1)]:
        doc.styles[n].font.name='Aptos'; doc.styles[n].font.size=Pt(sz)

def title(doc,a,b):
    p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(a)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(b); r.bold=True

def table(doc, rows, headers=('Item','Meaning / Value')):
    t=doc.add_table(rows=1, cols=2); t.style='Table Grid'; t.cell(0,0).text=headers[0]; t.cell(0,1).text=headers[1]; shade(t.cell(0,0),'D9EAF7'); shade(t.cell(0,1),'D9EAF7')
    for a,b in rows:
        c=t.add_row().cells; c[0].text=str(a); c[1].text=str(b)

def code(doc,s):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.12); r=p.add_run(textwrap.dedent(s).strip()); r.font.name='Consolas'; r.font.size=Pt(8.05); sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),'F3F4F6'); p._p.get_or_add_pPr().append(sh)

def note(doc,s):
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'; shade(t.cell(0,0),'FFF7D6'); p=t.cell(0,0).paragraphs[0]; r=p.add_run('Important: '); r.bold=True; p.add_run(s)

def bullets(doc, items):
    for x in items: doc.add_paragraph(x, style='List Bullet')

def sig(tid):
    return {'T16':'CategoryResponseDto createCategory(CategoryCreateRequestDto request)','T17':'CategoryResponseDto readCategory(Long id)','T18':'CategoryResponseDto updateCategory(Long id, CategoryUpdateRequestDto request)','T19':'CategoryResponseDto deleteCategory(Long id)','T20':'List<CategoryResponseDto> searchCategory(String text)'}[tid]

def empty_impl(tid,d):
    ret='return java.util.Collections.emptyList();' if tid=='T20' else 'return null;'
    return f'''@Service\npublic class {d['impl']} implements {d['service']} {{\n    @Override\n    public {sig(tid)} {{\n        {ret}\n    }}\n}}'''

def normal_impl(tid):
    return {
'T16':'''@Autowired private CategoryDao categoryDao;\n@Autowired private CategoryDtoDoMapper mapper;\n\n@Override\npublic CategoryResponseDto createCategory(CategoryCreateRequestDto request) {\n    CategoryDO dataObject = mapper.toDO(request);\n    CategoryDO saved = categoryDao.save(dataObject);\n    return mapper.toResponse(saved);\n}''',
'T17':'''@Autowired private CategoryDao categoryDao;\n@Autowired private CategoryDtoDoMapper mapper;\n\n@Override\npublic CategoryResponseDto readCategory(Long id) {\n    CategoryDO category = categoryDao.findById(id).orElse(null);\n    return mapper.toResponse(category);\n}''',
'T18':'''@Autowired private CategoryDao categoryDao;\n@Autowired private CategoryDtoDoMapper mapper;\n\n@Override\npublic CategoryResponseDto updateCategory(Long id, CategoryUpdateRequestDto request) {\n    CategoryDO category = categoryDao.findById(id).orElse(null);\n    mapper.applyUpdate(request, category);\n    return mapper.toResponse(categoryDao.save(category));\n}''',
'T19':'''@Autowired private CategoryDao categoryDao;\n@Autowired private CategoryDtoDoMapper mapper;\n\n@Override\npublic CategoryResponseDto deleteCategory(Long id) {\n    CategoryDO category = categoryDao.findById(id).orElse(null);\n    category.setActive(false);\n    return mapper.toResponse(categoryDao.save(category));\n}''',
'T20':'''@Autowired private CategoryDao categoryDao;\n@Autowired private CategoryDtoDoMapper mapper;\n\n@Override\npublic List<CategoryResponseDto> searchCategory(String text) {\n    String normalized = text.trim();\n    return categoryDao.search(normalized).stream().map(mapper::toResponse).toList();\n}'''}[tid]

def observe(tid):
    return {'T16':'service.createCategory(null);','T17':'when(categoryDao.findById(999L)).thenReturn(Optional.empty());\nservice.readCategory(999L);','T18':'when(categoryDao.findById(2L)).thenReturn(Optional.of(category(2L, "DATABASE", true)));\nservice.updateCategory(2L, null);','T19':'when(categoryDao.findById(999L)).thenReturn(Optional.empty());\nservice.deleteCategory(999L);','T20':'service.searchCategory(null);'}[tid]

def corrected(tid):
    return {
'T16':'''if (request == null) throw new IllegalArgumentException("Category request is required");\nString code = request.getCategoryCode().trim();\ncategoryDao.findByNormalizedBusinessKey(code).ifPresent(existing -> { throw new IllegalStateException("Category code already exists"); });\nrequest.setCategoryCode(code);\nrequest.setCategoryName(request.getCategoryName().trim());\nreturn mapper.toResponse(categoryDao.save(mapper.toDO(request)));''',
'T17':'''if (id == null) throw new IllegalArgumentException("Category ID is required");\nCategoryDO category = categoryDao.findCurrentById(id).orElseThrow(() -> new IllegalArgumentException("Category not found"));\nreturn mapper.toResponse(category);''',
'T18':'''if (id == null || request == null) throw new IllegalArgumentException("Category ID and request are required");\nCategoryDO category = categoryDao.findCurrentById(id).orElseThrow(() -> new IllegalArgumentException("Category not found"));\nString code = request.getCategoryCode().trim();\ncategoryDao.findByNormalizedBusinessKey(code).filter(other -> !other.getCategoryId().equals(id)).ifPresent(other -> { throw new IllegalStateException("Category code already exists"); });\nrequest.setCategoryCode(code); request.setCategoryName(request.getCategoryName().trim());\nmapper.applyUpdate(request, category);\nreturn mapper.toResponse(categoryDao.save(category));''',
'T19':'''if (id == null) throw new IllegalArgumentException("Category ID is required");\nCategoryDO category = categoryDao.findCurrentById(id).orElseThrow(() -> new IllegalArgumentException("Category not found"));\nif (categoryDao.countDependencies(id) > 0) throw new IllegalStateException("Category has active Books");\ncategory.setActive(false);\nreturn mapper.toResponse(categoryDao.save(category));''',
'T20':'''if (text == null) throw new IllegalArgumentException("Search text is required");\nString normalized = text.trim();\nif (normalized.isBlank()) throw new IllegalArgumentException("Search text is required");\nreturn categoryDao.searchIgnoreCase(normalized).stream().map(mapper::toResponse).toList();'''}[tid]

def integration(tid):
    return {'T16':'Create new SECURITY, then assert duplicate PROGRAMMING throws IllegalStateException.','T17':'Read ID 2 successfully, then assert ID 3 (inactive LEGACY) is not exposed.','T18':'Allow ID 1 to keep PROGRAMMING, then reject ID 2 changing to PROGRAMMING.','T19':'Deactivate unused NETWORK (ID 4), then reject PROGRAMMING (ID 1) because active Books depend on it.','T20':'Search programming and assert PROGRAMMING is returned; a valid no-match search returns an empty list.'}[tid]

def build_initial(tid,d):
    doc=Document(); setup(doc); title(doc,f'{tid} - {d["op"]} - Initial API Contract',f'{d["student"]} | {CONTRACT}')
    doc.add_heading('1. Fixed Assignment',1); table(doc,[('Track',tid),('Student',d['student']),('Operation',d['op']),('Service code',d['code']),('Endpoint',d['endpoint']),('Service interface',d['service']),('Method',sig(tid)),('Success responseCode','00'),('Success message',d['success'])])
    doc.add_heading('2. JSON Response Envelope',1); code(doc,f'''{{\n  "serviceCode": "{d['code']}",\n  "responseCode": "00",\n  "message": "{d['success']}",\n  "data": {{ }}\n}}''')
    doc.add_heading('3. Category Data',1); table(doc,[('categoryId','Database identifier'),('categoryCode','Business code such as PROGRAMMING'),('categoryName','Display name such as Programming'),('active','Current logical state')])
    doc.add_heading('4. Starting PostgreSQL Data',1); table(doc,[('1 / PROGRAMMING','Programming - ACTIVE'),('2 / DATABASE','Database Systems - ACTIVE'),('3 / LEGACY','Archived Category - INACTIVE'),('4 / NETWORK','Computer Networks - ACTIVE and unused by current Books')])
    doc.add_heading('5. Ownership',1); table(doc,[('Presenter','REST Controller, DTO/DO, Mapper, DAO base, Flyway, JSON envelope, Thymeleaf Controller/page framework'),('Student',f'{d["impl"]}, tests and {d["ui"]}')]); note(doc,'The Initial API Contract gives the successful starting behavior. Follow the Progressive Development Guide in sequence.')
    return doc

def build_progressive(tid,d):
    doc=Document(); setup(doc); title(doc,f'{d["student"]} - {d["op"]} Progressive Testing Guide',f'{tid} | Category Module | {CONTRACT}')
    doc.add_paragraph('Develop the service progressively. First prove that the method can be invoked. Then implement the normal successful path. Only after it works, try the next realistic condition, observe the result, improve the service, and verify the improvement with Unit Testing and real PostgreSQL Integration Testing.')
    doc.add_heading('1. Assignment and Request Flow',1); table(doc,[('Backend assignment',d['op']),('Service',d['service']),('Implementation',d['impl']),('Endpoint',d['endpoint']),('Service code',d['code']),('Frontend assignment',d['ui']),('REST Controller','Presenter-owned')]); table(doc,[('Browser / REST client','Sends HTTP'),('Spring Web + Jackson','Creates request DTO for POST/PUT'),('REST Controller','Calls fixed service'),('Service','Student business logic'),('Mapper','DTO to/from DO'),('DAO','JPA persistence'),('PostgreSQL','Real stored data')]); note(doc,'The service receives Java DTO values, not raw JSON.')
    doc.add_heading('2. Create the Empty Service',1); doc.add_paragraph('Unit Testing removes the complete Spring chain so the assigned service can be invoked independently. Mockito represents dependencies without starting the real database. This proves invocation before a large service is written.'); code(doc,empty_impl(tid,d)); code(doc,f'''@ExtendWith(MockitoExtension.class)\nclass {d['impl']}Test {{\n    @InjectMocks private {d['impl']} service;\n    @Test void shouldInvokeService() {{ /* call {sig(tid)} */ }}\n}}''')
    doc.add_heading('3. Add the Normal Successful Behavior',1); code(doc,normal_impl(tid)); doc.add_paragraph('Add one positive Mockito Unit Test for a valid request. Assert the returned Category values and verify the expected DAO call. Keep that test permanently as a regression test.')
    doc.add_heading('4. Try the Next Realistic Condition',1); doc.add_paragraph(d['discovery']); code(doc,f'''@Test\nvoid observeCurrentBehavior() {{\n{textwrap.indent(observe(tid), '    ')}\n}}'''); doc.add_paragraph('Read the exception type and the first stack-trace line in your service. A NullPointerException means a null reference was used as if it referred to an object. Understand the cause before changing the code.')
    doc.add_heading('5. Improve the Service',1); code(doc,corrected(tid)); doc.add_paragraph('Rerun both the new test and the earlier positive test. The correction is complete only when the successful behavior is still green.')
    doc.add_heading('6. JaCoCo Coverage',1); code(doc,'mvnw.cmd clean test'); table(doc,[('Green','Executed by tests'),('Yellow','Partly covered branch'),('Red','Not executed; there is still uncertainty')])
    doc.add_heading('7. Integration Testing',1); doc.add_paragraph('Mockito does not prove PostgreSQL comparison rules, active-state queries, persisted dependencies, Flyway data or JPA mappings. Integration Testing connects the real Spring components.'); table(doc,[('Unit Test','Service + mocks'),('Local Integration Test','Spring + Mapper + DAO + local PostgreSQL'),('Testcontainers','Spring + Mapper + DAO + temporary PostgreSQL')])
    doc.add_heading('8. JDBC Driver Setup',1); code(doc,'''spring.datasource.url=jdbc:postgresql://localhost:5432/library_full_stack\nspring.datasource.username=postgres\nspring.datasource.password=postgres\nspring.datasource.driver-class-name=org.postgresql.Driver'''); code(doc,'''<dependency>\n  <groupId>org.postgresql</groupId>\n  <artifactId>postgresql</artifactId>\n  <scope>runtime</scope>\n</dependency>\n\n<!-- MySQL learning reference only - keep commented\n<dependency>\n  <groupId>com.mysql</groupId>\n  <artifactId>mysql-connector-j</artifactId>\n  <scope>runtime</scope>\n</dependency>\n-->'''); code(doc,'''# MySQL reference only\n# spring.datasource.url=jdbc:mysql://localhost:3306/library_full_stack\n# spring.datasource.driver-class-name=com.mysql.cj.jdbc.Driver''')
    doc.add_heading('9. Inspect the Training Data in DBeaver',1); code(doc,'''SELECT category_id, category_code, category_name, is_active FROM tbl_category ORDER BY category_id;\nSELECT book_id, isbn, title, fk_category, is_active FROM tbl_book ORDER BY book_id;'''); table(doc,[('Category 1','PROGRAMMING - ACTIVE and used by active Books'),('Category 2','DATABASE - ACTIVE'),('Category 3','LEGACY - INACTIVE'),('Category 4','NETWORK - ACTIVE and unused by current Books')])
    doc.add_heading('10. Main Integration Test',1); doc.add_paragraph(d['integration']); code(doc,integration(tid));
    if tid=='T20': code(doc,'''@Transactional(readOnly = true)\npublic List<CategoryDO> searchIgnoreCase(String text) {\n    String pattern = "%" + text.toLowerCase() + "%";\n    return entityManager.createQuery(\n        "select c from CategoryDO c where c.active = true and " +\n        "(lower(c.categoryCode) like :pattern or lower(c.categoryName) like :pattern) " +\n        "order by c.categoryId", CategoryDO.class)\n        .setParameter("pattern", pattern)\n        .getResultList();\n}''')
    doc.add_heading('11. PostgreSQL Testcontainers',1); code(doc,'''@Container\n@ServiceConnection\nstatic PostgreSQLContainer postgres =\n    new PostgreSQLContainer(DockerImageName.parse("postgres:18"));'''); doc.add_paragraph('Add @Testcontainers to the Integration Test class and keep the same test cases. Flyway rebuilds the deterministic data in the temporary database.')
    doc.add_heading('12. Final Regression',1); code(doc,'mvnw.cmd clean test'); bullets(doc,['Focused Unit Test green','Original successful test still green','Corrected condition test green','Local PostgreSQL Integration Test green','Testcontainers test green','JaCoCo reviewed'])
    doc.add_heading('13. Frontend Assignment',1); doc.add_paragraph(f'{d["student"]} implements the {d["ui"]}. First prove main.js and the button handler, then build one working JavaScript file, test Presenter STUB/support, and only then refactor into main.js, API, form and view modules.'); code(doc,'''<script type="module" src="/src/main.js"></script>\n\nconsole.log("main.js loaded");\ndocument.querySelector("#actionButton").addEventListener("click", () => console.log("action button reached JavaScript"));'''); code(doc,'''if (response.status === 404) { showMessage("NOT Found"); return; }\nif (!response.ok) { showMessage("Error Encountered, Please contact Administrator"); return; }\nconst body = await response.json();\nshowMessage(body.message);'''); bullets(doc,['single-file page works','Presenter STUB/support works','HTTP handling works','page refactored only after working','partner implementation works with unchanged /rest URL'])
    return doc

def build_updated(tid,d):
    doc=Document(); setup(doc); title(doc,f'{tid} - {d["op"]} - Updated API Contract',f'{CONTRACT} | Category Module'); doc.add_heading('1. Fixed Endpoint',1); table(doc,[('Track',tid),('Service code',d['code']),('Endpoint',d['endpoint']),('Controlled application result','HTTP 200')]); doc.add_heading('2. JSON Envelope',1); code(doc,'''{\n  "serviceCode": "XX",\n  "responseCode": "00",\n  "message": "...",\n  "data": {}\n}'''); codes=[('00',d['success'])]
    if tid=='T16': codes += [('01','Invalid / missing Category request'),('03','Category code already exists')]
    if tid=='T17': codes += [('01','Category ID is required'),('02','Category not found / inactive row not exposed')]
    if tid=='T18': codes += [('01','Invalid Category ID/request'),('02','Category not found'),('03','Category code belongs to another Category')]
    if tid=='T19': codes += [('01','Category ID is required'),('02','Category not found'),('05','Category has active Book dependencies')]
    if tid=='T20': codes += [('01','Search text is required or blank')]
    doc.add_heading('3. Response Codes',1); table(doc,codes,('responseCode','Meaning')); doc.add_heading('4. Frontend Transport Errors',1); table(doc,[('HTTP 404','NOT Found'),('Other non-200','Error Encountered, Please contact Administrator')]); note(doc,'A valid Search with zero matches returns responseCode 00 and an empty data list.'); return doc

for tid,d in TRACKS.items():
    folder=ROOT/f"{tid}_{d['op'].replace(' ','_')}"; folder.mkdir(parents=True, exist_ok=True)
    build_initial(tid,d).save(folder/f'{tid}_01_Initial_API_Contract.docx')
    build_progressive(tid,d).save(folder/f'{tid}_02_Progressive_Development_Guide.docx')
    build_updated(tid,d).save(folder/f'{tid}_03_Updated_API_Contract.docx')

print('Generated T16-T20 Category guide set')
