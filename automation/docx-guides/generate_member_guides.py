from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import textwrap

CONTRACT = "LIBRARY-FULLSTACK-V1"
OUT = Path("Documents/Student-Guides/Member")

TRACKS = {
    "T06": dict(student="Student 06", operation="Create Member", service="CreateMemberService", impl="CreateMemberServiceImpl", code="06", endpoint="POST /rest/members", success="Member Created Successfully", frontend="Update Member UI", unit="Call createMember(null) after the successful create path works.", integration="The database already contains REG-CSE-001. Try to create another Member using the same registration number."),
    "T07": dict(student="Student 07", operation="Read Member", service="ReadMemberService", impl="ReadMemberServiceImpl", code="07", endpoint="GET /rest/members/{id}", success="Member Read Successfully", frontend="Deactivate Member UI", unit="Request Member ID 999 after the successful read path works.", integration="Member ID 3 physically exists as REG-CIV-001 but is inactive. Normal Read must not expose it."),
    "T08": dict(student="Student 08", operation="Update Member", service="UpdateMemberService", impl="UpdateMemberServiceImpl", code="08", endpoint="PUT /rest/members/{id}", success="Member Updated Successfully", frontend="Search Member UI", unit="Call updateMember(2L, null) after the successful update path works.", integration="Member ID 1 owns REG-CSE-001. Try to update Member ID 2 from REG-ECE-001 to REG-CSE-001. Duplicate checking must exclude the current row."),
    "T09": dict(student="Student 09", operation="Deactivate Member", service="DeleteMemberService", impl="DeleteMemberServiceImpl", code="09", endpoint="DELETE /rest/members/{id}", success="Member Deactivated Successfully", frontend="Create Member UI", unit="Request Member ID 999 after logical deactivation works for a valid Member.", integration="Member ID 1 has an ACTIVE membership used by ACTIVE Book Issue ISS-0001. It must not be deactivated while that persisted dependency exists."),
    "T10": dict(student="Student 10", operation="Search Member", service="SearchMemberService", impl="SearchMemberServiceImpl", code="10", endpoint="GET /rest/members/search?text={text}", success="Member Search Completed Successfully", frontend="Member List / Search Results UI", unit="Call searchMember(null) after a normal search works. The first implementation trims the text.", integration="Search for reg-cse-001 and arun while PostgreSQL stores REG-CSE-001 and Arun Kumar. The final search must be case-independent."),
}

METHODS = {
    "T06": "MemberResponseDto createMember(MemberCreateRequestDto request)",
    "T07": "MemberResponseDto readMember(Long id)",
    "T08": "MemberResponseDto updateMember(Long id, MemberUpdateRequestDto request)",
    "T09": "MemberResponseDto deleteMember(Long id)",
    "T10": "List<MemberResponseDto> searchMember(String text)",
}

NORMAL = {
"T06": '''@Autowired private MemberDao memberDao;
@Autowired private MemberDtoDoMapper mapper;

@Override
public MemberResponseDto createMember(MemberCreateRequestDto request) {
    MemberDO dataObject = mapper.toDO(request);
    MemberDO saved = memberDao.save(dataObject);
    return mapper.toResponse(saved);
}''',
"T07": '''@Autowired private MemberDao memberDao;
@Autowired private MemberDtoDoMapper mapper;

@Override
public MemberResponseDto readMember(Long id) {
    MemberDO member = memberDao.findById(id).orElse(null);
    return mapper.toResponse(member);
}''',
"T08": '''@Autowired private MemberDao memberDao;
@Autowired private MemberDtoDoMapper mapper;

@Override
public MemberResponseDto updateMember(Long id, MemberUpdateRequestDto request) {
    MemberDO member = memberDao.findById(id).orElse(null);
    mapper.applyUpdate(request, member);
    return mapper.toResponse(memberDao.save(member));
}''',
"T09": '''@Autowired private MemberDao memberDao;
@Autowired private MemberDtoDoMapper mapper;

@Override
public MemberResponseDto deleteMember(Long id) {
    MemberDO member = memberDao.findById(id).orElse(null);
    member.setActive(false);
    return mapper.toResponse(memberDao.save(member));
}''',
"T10": '''@Autowired private MemberDao memberDao;
@Autowired private MemberDtoDoMapper mapper;

@Override
public List<MemberResponseDto> searchMember(String text) {
    String normalized = text.trim();
    return memberDao.search(normalized).stream()
            .map(mapper::toResponse)
            .toList();
}'''
}

POSITIVE = {
"T06": '''@Test
void shouldCreateMember() {
    MemberCreateRequestDto request = createRequest(
            "REG-NEW-001", "New Student", "new@example.edu", 1L);
    MemberDO dataObject = member(null, "REG-NEW-001", true);
    MemberDO saved = member(100L, "REG-NEW-001", true);
    when(mapper.toDO(request)).thenReturn(dataObject);
    when(memberDao.save(dataObject)).thenReturn(saved);
    when(mapper.toResponse(saved)).thenReturn(
            response(100L, "REG-NEW-001", true));

    MemberResponseDto result = service.createMember(request);

    assertEquals("REG-NEW-001", result.getRegistrationNumber());
    verify(memberDao).save(dataObject);
}''',
"T07": '''@Test
void shouldReadMember() {
    MemberDO existing = member(2L, "REG-ECE-001", true);
    when(memberDao.findById(2L)).thenReturn(Optional.of(existing));
    when(mapper.toResponse(existing)).thenReturn(
            response(2L, "REG-ECE-001", true));

    MemberResponseDto result = service.readMember(2L);

    assertEquals("REG-ECE-001", result.getRegistrationNumber());
    verify(memberDao).findById(2L);
}''',
"T08": '''@Test
void shouldUpdateMember() {
    MemberDO existing = member(2L, "REG-ECE-001", true);
    MemberUpdateRequestDto request = updateRequest(
            "REG-ECE-001", "Priya Devi", "priya@example.edu", 2L);
    when(memberDao.findById(2L)).thenReturn(Optional.of(existing));
    when(memberDao.save(existing)).thenReturn(existing);
    when(mapper.toResponse(existing)).thenReturn(
            response(2L, "REG-ECE-001", true));

    MemberResponseDto result = service.updateMember(2L, request);

    assertEquals("REG-ECE-001", result.getRegistrationNumber());
    verify(mapper).applyUpdate(request, existing);
    verify(memberDao).save(existing);
}''',
"T09": '''@Test
void shouldDeactivateMember() {
    MemberDO existing = member(4L, "REG-MECH-001", true);
    when(memberDao.findById(4L)).thenReturn(Optional.of(existing));
    when(memberDao.save(existing)).thenReturn(existing);
    when(mapper.toResponse(existing)).thenReturn(
            response(4L, "REG-MECH-001", false));

    MemberResponseDto result = service.deleteMember(4L);

    assertFalse(existing.getActive());
    assertFalse(result.getActive());
    verify(memberDao).save(existing);
}''',
"T10": '''@Test
void shouldSearchMembers() {
    MemberDO arun = member(1L, "REG-CSE-001", true);
    when(memberDao.search("Arun")).thenReturn(List.of(arun));
    when(mapper.toResponse(arun)).thenReturn(
            response(1L, "REG-CSE-001", true));

    List<MemberResponseDto> result = service.searchMember(" Arun ");

    assertEquals(1, result.size());
    assertEquals("REG-CSE-001", result.get(0).getRegistrationNumber());
}'''
}

OBSERVE = {
"T06": '''@Test
void createWithNullRequestObserveCurrentBehavior() {
    service.createMember(null);
}''',
"T07": '''@Test
void readUnknownMemberObserveCurrentBehavior() {
    when(memberDao.findById(999L)).thenReturn(Optional.empty());
    service.readMember(999L);
}''',
"T08": '''@Test
void updateWithNullRequestObserveCurrentBehavior() {
    when(memberDao.findById(2L)).thenReturn(
            Optional.of(member(2L, "REG-ECE-001", true)));
    service.updateMember(2L, null);
}''',
"T09": '''@Test
void deactivateUnknownMemberObserveCurrentBehavior() {
    when(memberDao.findById(999L)).thenReturn(Optional.empty());
    service.deleteMember(999L);
}''',
"T10": '''@Test
void searchWithNullTextObserveCurrentBehavior() {
    service.searchMember(null);
}'''
}

CORRECTED = {
"T06": '''@Override
public MemberResponseDto createMember(MemberCreateRequestDto request) {
    if (request == null) {
        throw new IllegalArgumentException("Member request is required");
    }
    String registrationNumber = request.getRegistrationNumber().trim();
    memberDao.findByNormalizedBusinessKey(registrationNumber)
            .ifPresent(existing -> {
                throw new IllegalStateException(
                        "Registration number already exists");
            });
    request.setRegistrationNumber(registrationNumber);
    request.setFullName(request.getFullName().trim());
    request.setEmail(request.getEmail().trim());
    return mapper.toResponse(memberDao.save(mapper.toDO(request)));
}''',
"T07": '''@Override
public MemberResponseDto readMember(Long id) {
    if (id == null) {
        throw new IllegalArgumentException("Member ID is required");
    }
    MemberDO member = memberDao.findCurrentById(id)
            .orElseThrow(() ->
                    new IllegalArgumentException("Member not found"));
    return mapper.toResponse(member);
}''',
"T08": '''@Override
public MemberResponseDto updateMember(
        Long id, MemberUpdateRequestDto request) {
    if (id == null || request == null) {
        throw new IllegalArgumentException(
                "Member ID and request are required");
    }
    MemberDO member = memberDao.findCurrentById(id)
            .orElseThrow(() ->
                    new IllegalArgumentException("Member not found"));
    String registrationNumber = request.getRegistrationNumber().trim();
    memberDao.findByNormalizedBusinessKey(registrationNumber)
            .filter(other -> !other.getMemberId().equals(id))
            .ifPresent(other -> {
                throw new IllegalStateException(
                        "Registration number already exists");
            });
    request.setRegistrationNumber(registrationNumber);
    request.setFullName(request.getFullName().trim());
    request.setEmail(request.getEmail().trim());
    mapper.applyUpdate(request, member);
    return mapper.toResponse(memberDao.save(member));
}''',
"T09": '''@Override
public MemberResponseDto deleteMember(Long id) {
    if (id == null) {
        throw new IllegalArgumentException("Member ID is required");
    }
    MemberDO member = memberDao.findCurrentById(id)
            .orElseThrow(() ->
                    new IllegalArgumentException("Member not found"));
    if (memberDao.countDependencies(id) > 0) {
        throw new IllegalStateException(
                "Member has an active Book Issue");
    }
    member.setActive(false);
    return mapper.toResponse(memberDao.save(member));
}''',
"T10": '''@Override
public List<MemberResponseDto> searchMember(String text) {
    if (text == null) {
        throw new IllegalArgumentException("Search text is required");
    }
    String normalized = text.trim();
    if (normalized.isBlank()) {
        throw new IllegalArgumentException("Search text is required");
    }
    return memberDao.searchIgnoreCase(normalized).stream()
            .map(mapper::toResponse)
            .toList();
}'''
}

INTEGRATION = {
"T06": '''@SpringBootTest
class CreateMemberLocalIntegrationTest {
    @Autowired private CreateMemberService service;

    @Test
    void shouldCreateNewRegistrationNumber() {
        MemberResponseDto result = service.createMember(
                createRequest("REG-NEW-001", "New Student",
                        "new@example.edu", 1L));
        assertEquals("REG-NEW-001", result.getRegistrationNumber());
    }

    @Test
    void shouldRejectPersistedDuplicateRegistrationNumber() {
        assertThrows(IllegalStateException.class, () ->
                service.createMember(createRequest(
                        "REG-CSE-001", "Another Student",
                        "another@example.edu", 1L)));
    }
}''',
"T07": '''@SpringBootTest
class ReadMemberLocalIntegrationTest {
    @Autowired private ReadMemberService service;

    @Test
    void shouldReadActiveMember() {
        MemberResponseDto result = service.readMember(2L);
        assertEquals("REG-ECE-001", result.getRegistrationNumber());
        assertTrue(result.getActive());
    }

    @Test
    void normalReadShouldNotExposeInactiveMember() {
        assertThrows(IllegalArgumentException.class,
                () -> service.readMember(3L));
    }
}''',
"T08": '''@SpringBootTest
class UpdateMemberLocalIntegrationTest {
    @Autowired private UpdateMemberService service;

    @Test
    void shouldAllowCurrentMemberToKeepOwnRegistrationNumber() {
        MemberResponseDto result = service.updateMember(1L,
                updateRequest("REG-CSE-001", "Arun Kumar",
                        "arun@example.edu", 1L));
        assertEquals("REG-CSE-001", result.getRegistrationNumber());
    }

    @Test
    void shouldRejectRegistrationNumberOwnedByAnotherMember() {
        assertThrows(IllegalStateException.class, () ->
                service.updateMember(2L,
                        updateRequest("REG-CSE-001", "Priya Devi",
                                "priya@example.edu", 2L)));
    }
}''',
"T09": '''@SpringBootTest
class DeleteMemberLocalIntegrationTest {
    @Autowired private DeleteMemberService service;

    @Test
    void shouldDeactivateMemberWithoutActiveIssueDependency() {
        assertFalse(service.deleteMember(4L).getActive());
    }

    @Test
    void shouldNotDeactivateMemberWithActiveBookIssue() {
        assertThrows(IllegalStateException.class,
                () -> service.deleteMember(1L));
    }
}''',
"T10": '''@SpringBootTest
class SearchMemberLocalIntegrationTest {
    @Autowired private SearchMemberService service;

    @Test
    void shouldFindUppercaseRegistrationUsingLowercaseSearch() {
        List<MemberResponseDto> result = service.searchMember("reg-cse-001");
        assertTrue(result.stream().anyMatch(m ->
                "REG-CSE-001".equals(m.getRegistrationNumber())));
    }

    @Test
    void shouldFindMixedCaseNameUsingLowercaseSearch() {
        List<MemberResponseDto> result = service.searchMember("arun");
        assertTrue(result.stream().anyMatch(m ->
                "Arun Kumar".equals(m.getFullName())));
    }
}'''
}

UPDATED = {
    "T06": [("00", "Member Created Successfully"), ("01", "Invalid / missing Member request"), ("03", "Registration number already exists")],
    "T07": [("00", "Member Read Successfully"), ("01", "Member ID is required"), ("02", "Member not found / inactive record is not exposed")],
    "T08": [("00", "Member Updated Successfully"), ("01", "Invalid Member ID/request"), ("02", "Member not found"), ("03", "Registration number belongs to another Member")],
    "T09": [("00", "Member Deactivated Successfully"), ("01", "Member ID is required"), ("02", "Member not found"), ("05", "Member has an active dependent Book Issue")],
    "T10": [("00", "Member Search Completed Successfully"), ("01", "Search text is required or blank")],
}


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def setup(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(10.2)
    for name, size in (("Title", 21), ("Heading 1", 15), ("Heading 2", 12.8), ("Heading 3", 11.2)):
        doc.styles[name].font.name = "Aptos"
        doc.styles[name].font.size = Pt(size)


def title(doc, heading, subtitle):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(heading)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.bold = True


def table(doc, rows, headers=("Item", "Meaning / Value")):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = headers[0]
    t.cell(0, 1).text = headers[1]
    shade(t.cell(0, 0), "D9EAF7")
    shade(t.cell(0, 1), "D9EAF7")
    for left, right in rows:
        cells = t.add_row().cells
        cells[0].text = str(left)
        cells[1].text = str(right)


def code(doc, source):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.13)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(textwrap.dedent(source).strip())
    r.font.name = "Consolas"
    r.font.size = Pt(8.15)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F4F6")
    p._p.get_or_add_pPr().append(shd)


def note(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    shade(t.cell(0, 0), "FFF7D6")
    t.cell(0, 0).text = text


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_initial(track, d):
    doc = Document()
    setup(doc)
    title(doc, f"{track} - {d['operation']} - Initial API Contract", f"{d['student']} | {CONTRACT}")
    doc.add_heading("1. Assignment", level=1)
    table(doc, [("Student", d["student"]), ("Track", track), ("Operation", d["operation"]), ("Service Code", d["code"]), ("Endpoint", d["endpoint"]), ("Service Interface", d["service"]), ("Method", METHODS[track]), ("Success Response Code", "00"), ("Success Message", d["success"])])
    doc.add_heading("2. Fixed JSON Response Envelope", level=1)
    code(doc, f'''{{\n  "serviceCode": "{d['code']}",\n  "responseCode": "00",\n  "message": "{d['success']}",\n  "data": {{ }}\n}}''')
    doc.add_heading("3. Member DTO Fields", level=1)
    table(doc, [("registrationNumber", "Member business key, for example REG-CSE-001"), ("fullName", "Member full name"), ("email", "Member email address"), ("departmentId", "Department database identifier"), ("active", "Current/inactive state returned in MemberResponseDto")])
    doc.add_heading("4. Training Data", level=1)
    table(doc, [("1", "REG-CSE-001 / Arun Kumar / ACTIVE"), ("2", "REG-ECE-001 / Priya Devi / ACTIVE"), ("3", "REG-CIV-001 / Legacy Member / INACTIVE"), ("4", "REG-MECH-001 / Meena S / ACTIVE")])
    doc.add_heading("5. Ownership", level=1)
    table(doc, [("Presenter", "REST Controller, DTOs, DOs, Mapper, DAO base, Flyway data, common response, Thymeleaf shell"), ("Student", f"{d['impl']}, track tests, {d['frontend']}")])
    note(doc, "Only the successful API behavior is supplied at the beginning. Follow the Progressive Development Guide in order.")
    return doc


def build_progressive(track, d):
    doc = Document()
    setup(doc)
    title(doc, f"{d['student']} - {d['operation']} Progressive Testing Guide", f"{track} | Member Module | {CONTRACT}")
    doc.add_paragraph("Develop the service in small observable steps. Prove invocation, implement the normal successful behavior, try the next realistic condition before changing the code, study the result, improve the implementation, and then verify the completed behavior against real PostgreSQL.")

    doc.add_heading("1. Assignment and Request Flow", level=1)
    table(doc, [("Track", track), ("Backend", d["operation"]), ("Service", d["service"]), ("Method", METHODS[track]), ("REST Endpoint", d["endpoint"]), ("Service Code", d["code"]), ("Frontend", d["frontend"]), ("REST Controller", "Presenter-owned")])
    table(doc, [("Browser / client", "Sends HTTP request"), ("Spring Web + Jackson", "Converts request JSON into the supplied DTO"), ("REST Controller", "Calls the fixed service interface"), ("Service", "Student business logic"), ("Mapper", "DTO ↔ DO conversion"), ("DAO", "JPA persistence"), ("PostgreSQL", "Real persisted data")])
    note(doc, "The Service receives Java DTO values. It does not parse raw JSON; Jackson has already performed that conversion before the Controller calls the Service.")

    doc.add_heading("2. Create the Empty Service", level=1)
    doc.add_paragraph("A full Spring request passes through many components. If a large service is written before it is ever invoked, a later failure can be difficult to locate. The empty service removes this uncertainty. Unit Testing can call the service directly, similar to calling an ordinary Java object from a small main method. Mockito supplies controlled mock dependencies; the real database is introduced later during Integration Testing.")
    ret = "return java.util.Collections.emptyList();" if track == "T10" else "return null;"
    code(doc, f'''@Service\npublic class {d['impl']} implements {d['service']} {{\n    @Override\n    public {METHODS[track]} {{\n        {ret}\n    }}\n}}''')

    doc.add_heading("3. Invoke the Empty Service with JUnit", level=1)
    code(doc, f'''@ExtendWith(MockitoExtension.class)\nclass {d['impl']}Test {{\n    @InjectMocks\n    private {d['impl']} service;\n\n    @Test\n    void shouldInvokeService() {{\n        // Invoke {METHODS[track]} here.\n        // At this stage only invocation is being proved.\n    }}\n}}''')
    bullets(doc, ["Open the test class in Eclipse.", "Right-click the test method.", "Choose Run As → JUnit Test.", "Confirm that the JUnit view is green before adding business logic."])

    doc.add_heading(f"4. Implement the Normal {d['operation']} Behavior", level=1)
    code(doc, NORMAL[track])
    doc.add_paragraph("The first implementation is intentionally ordinary: it implements the valid business path that the Initial API Contract describes. Do not add every possible validation before the valid path has been proved.")

    doc.add_heading("5. Successful Unit Test", level=1)
    code(doc, POSITIVE[track])
    doc.add_paragraph("Keep this test permanently. It becomes a regression test. Later improvements are complete only when the original successful behavior still passes.")

    doc.add_heading("6. Try the Next Realistic Condition", level=1)
    doc.add_paragraph(d["unit"])
    code(doc, OBSERVE[track])
    doc.add_paragraph("Run the test before changing the service. In Eclipse, read the exception type and then the first stack-trace line belonging to the service implementation. A NullPointerException means a null reference was used as though an object existed. The application condition should be handled before that value is dereferenced.")

    doc.add_heading("7. Improve the Service", level=1)
    code(doc, CORRECTED[track])
    doc.add_paragraph("Rerun the new test and the earlier successful test. Do not remove the successful test after making the correction.")

    doc.add_heading("8. JaCoCo Coverage", level=1)
    code(doc, "mvnw.cmd clean test")
    table(doc, [("Green", "Executed by tests"), ("Yellow", "Partially covered branch/condition"), ("Red", "Not executed; still contains testing uncertainty")])
    doc.add_paragraph("Open target/site/jacoco/index.html and inspect the Member service implementation. Coverage shows what ran; it does not replace assertions or realistic test data.")

    doc.add_heading("9. Why Integration Testing is Required", level=1)
    table(doc, [("Unit Test", "Service + Mockito mocks"), ("Local Integration Test", "Spring + real Mapper + real DAO + local PostgreSQL"), ("Testcontainers", "Spring + real Mapper + real DAO + temporary PostgreSQL")])
    doc.add_paragraph(d["integration"])

    doc.add_heading("10. PostgreSQL JDBC Configuration", level=1)
    code(doc, '''spring.datasource.url=jdbc:postgresql://localhost:5432/library_full_stack\nspring.datasource.username=postgres\nspring.datasource.password=postgres\nspring.datasource.driver-class-name=org.postgresql.Driver''')
    code(doc, '''<dependency>\n    <groupId>org.postgresql</groupId>\n    <artifactId>postgresql</artifactId>\n    <scope>runtime</scope>\n</dependency>\n\n<!-- MySQL learning reference only - keep commented\n<dependency>\n    <groupId>com.mysql</groupId>\n    <artifactId>mysql-connector-j</artifactId>\n    <scope>runtime</scope>\n</dependency>\n-->''')
    code(doc, '''# MySQL learning reference only - DO NOT activate\n# spring.datasource.url=jdbc:mysql://localhost:3306/library_full_stack\n# spring.datasource.username=root\n# spring.datasource.password=your-password\n# spring.datasource.driver-class-name=com.mysql.cj.jdbc.Driver''')

    doc.add_heading("11. Inspect Training Data in DBeaver", level=1)
    code(doc, '''SELECT member_id, registration_number, full_name, email, fk_department, is_active\nFROM tbl_member\nORDER BY member_id;\n\nSELECT i.book_issue_id, i.issue_number, i.status, ms.member_id\nFROM tbl_book_issue i\nJOIN tbl_membership ms ON ms.membership_id = i.fk_membership\nORDER BY i.book_issue_id;''')
    table(doc, [("Member 1", "REG-CSE-001 / Arun Kumar / ACTIVE"), ("Member 2", "REG-ECE-001 / Priya Devi / ACTIVE"), ("Member 3", "REG-CIV-001 / Legacy Member / INACTIVE"), ("Member 4", "REG-MECH-001 / Meena S / ACTIVE"), ("Issue dependency", "Member 1 membership is used by ACTIVE issue ISS-0001")])

    doc.add_heading("12. Complete Local Integration Test", level=1)
    code(doc, INTEGRATION[track])
    if track == "T10":
        doc.add_heading("12.1 Correct the PostgreSQL Search", level=2)
        code(doc, '''@Transactional(readOnly = true)\npublic List<MemberDO> searchIgnoreCase(String text) {\n    String pattern = "%" + text.toLowerCase() + "%";\n    return entityManager.createQuery(\n            "select m from MemberDO m where m.active = true and " +\n            "(lower(m.registrationNumber) like :pattern or lower(m.fullName) like :pattern) " +\n            "order by m.memberId", MemberDO.class)\n            .setParameter("pattern", pattern)\n            .getResultList();\n}''')

    doc.add_heading("13. PostgreSQL Testcontainers", level=1)
    code(doc, '''@Container\n@ServiceConnection\nstatic PostgreSQLContainer postgres =\n        new PostgreSQLContainer(DockerImageName.parse("postgres:18"));''')
    doc.add_paragraph("Add @Testcontainers to the integration-test class and keep the same test methods. Spring Boot receives the JDBC connection from the PostgreSQL container, and Flyway rebuilds the deterministic schema/data in that temporary database.")

    doc.add_heading("14. Final Regression", level=1)
    code(doc, "mvnw.cmd clean test")
    bullets(doc, ["Focused JUnit test is green.", "Original successful test is still green.", "Local PostgreSQL integration test is green.", "PostgreSQL Testcontainers test is green.", "JaCoCo report has been reviewed."])

    doc.add_heading("15. Frontend Assignment", level=1)
    doc.add_paragraph(f"{d['student']} implements the {d['frontend']}. First build a complete working page in one JavaScript file. Prove that main.js runs and the button handler is reached before making the REST request. After that, refactor into main.js, an API module, a form module and a view module.")
    code(doc, '''<script type="module" src="/src/main.js"></script>\n\nconsole.log("main.js loaded");\ndocument.querySelector("#actionButton").addEventListener("click", () => {\n    console.log("action button reached JavaScript");\n});''')
    doc.add_heading("15.1 HTTP Transport Handling", level=2)
    code(doc, '''const response = await fetch(url, options);\n\nif (response.status === 404) {\n  showMessage("NOT Found");\n  return;\n}\n\nif (!response.ok) {\n  showMessage("Error Encountered, Please contact Administrator");\n  return;\n}\n\nconst body = await response.json();\nshowMessage(body.message);''')
    doc.add_heading("15.2 STUB to Partner Integration", level=2)
    doc.add_paragraph("Test the page first against Presenter STUB/support. When the real partner service is integrated, keep the same /rest URL. The HTTP contract remains stable while the backend implementation changes.")

    doc.add_heading("16. Completion Checklist", level=1)
    bullets(doc, ["Initial contract understood.", "Empty service invocation proved.", "Successful path implemented and tested.", "Next realistic condition executed before correction.", "Observed behavior understood.", "Service corrected and regression-tested.", "JaCoCo reviewed.", "Local PostgreSQL test completed.", "PostgreSQL Testcontainers test completed.", "Frontend proved against Presenter STUB/support.", "Frontend verified against the real partner implementation without changing the URL."])
    return doc


def build_updated(track, d):
    doc = Document()
    setup(doc)
    title(doc, f"{track} - {d['operation']} - Updated API Contract", f"{CONTRACT} | Released after progressive development")
    doc.add_heading("1. Fixed Endpoint", level=1)
    table(doc, [("Track", track), ("Service Code", d["code"]), ("Endpoint", d["endpoint"]), ("Controlled application results", "HTTP 200")])
    doc.add_heading("2. Response Envelope", level=1)
    code(doc, '''{\n  "serviceCode": "XX",\n  "responseCode": "00",\n  "message": "...",\n  "data": {}\n}''')
    doc.add_heading("3. Application Response Codes", level=1)
    table(doc, UPDATED[track], ("responseCode", "Meaning"))
    doc.add_heading("4. Frontend Transport Handling", level=1)
    table(doc, [("HTTP 404", "NOT Found"), ("Other non-200", "Error Encountered, Please contact Administrator")])
    note(doc, "A valid Search with no matching Member is still successful: responseCode 00 with an empty data list.")
    return doc


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    for track, d in TRACKS.items():
        folder = OUT / f"{track}_{d['operation'].replace(' ', '_')}"
        folder.mkdir(parents=True, exist_ok=True)
        build_initial(track, d).save(folder / f"{track}_01_Initial_API_Contract.docx")
        build_progressive(track, d).save(folder / f"{track}_02_Progressive_Development_Guide.docx")
        build_updated(track, d).save(folder / f"{track}_03_Updated_API_Contract.docx")
    print("Generated 15 Member DOCX guides in", OUT)


if __name__ == "__main__":
    main()
