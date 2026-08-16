from pathlib import Path

from docx import Document


ROOT = Path("Documents/Student-Guides/Department/T04_Deactivate_Department")
PROGRESSIVE = ROOT / "T04_02_Progressive_Development_Guide.docx"
INITIAL = ROOT / "T04_01_Initial_API_Contract.docx"


def replace_paragraph(paragraph):
    text = paragraph.text
    updated = text

    # Earlier material used existing Departments 2/4 as convenient successful
    # examples.  In the final aggregate seed those Departments have active
    # Members.  The audited success row is Department 5 / IT.
    updated = updated.replace("service.deleteDepartment(2L)", "service.deleteDepartment(5L)")
    updated = updated.replace("service.deleteDepartment(4L)", "service.deleteDepartment(5L)")

    lowered = updated.lower()
    if "department id 2" in lowered and (
            "deactivat" in lowered or "active member" in lowered or "success" in lowered):
        updated = updated.replace("Department ID 2", "Department ID 5")
    if "department id 4" in lowered and (
            "deactivat" in lowered or "active member" in lowered or "success" in lowered):
        updated = updated.replace("Department ID 4", "Department ID 5")

    if updated == text:
        return False

    paragraph.text = updated
    return True


def replace_everywhere(doc):
    changed = False
    for paragraph in doc.paragraphs:
        changed = replace_paragraph(paragraph) or changed
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changed = replace_paragraph(paragraph) or changed
    return changed


def add_audit_note(doc):
    marker = "Final audited T04 training row"
    if any(marker in p.text for p in doc.paragraphs):
        return False

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        "Final audited T04 training row: use Department ID 5 / IT "
        "(Information Technology) for the successful Deactivate integration "
        "case because it has no active Member. Department ID 1 / CSE remains "
        "the persisted dependency-block case."
    )
    run.bold = True
    return True


def add_training_row(doc):
    for table in doc.tables:
        texts = [cell.text for row in table.rows for cell in row.cells]
        joined = " | ".join(texts)
        if "CSE" not in joined or "MECH" not in joined:
            continue
        if "Information Technology" in joined:
            return False
        row = table.add_row().cells
        if len(row) >= 2:
            row[0].text = "5 / IT"
            row[1].text = "Information Technology - ACTIVE; no active Member; T04 successful Deactivate row"
            return True
    return False


def main():
    initial = Document(INITIAL)
    initial_changed = add_training_row(initial)
    if initial_changed:
        initial.save(INITIAL)

    progressive = Document(PROGRESSIVE)
    progressive_changed = replace_everywhere(progressive)
    progressive_changed = add_audit_note(progressive) or progressive_changed
    if progressive_changed:
        progressive.save(PROGRESSIVE)

    print("Aligned T04 guide to final audited Department ID 5 success row.")


if __name__ == "__main__":
    main()
